from __future__ import annotations

import io
import json
import math
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from datetime import datetime
from typing import Any, Callable

import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, PageBreak
from scipy.stats import norm
from sqlalchemy import create_engine, inspect


APP_NAME = "BioSize Clinical"
APP_VERSION = "1.2.0"
MIN_POWER = 0.80

OUTCOME_TYPE_LABELS: dict[str, str] = {
    "continuous": "Continuo",
    "binary": "Dicotómico/binario",
    "paired_binary": "Dicotómico apareado",
    "time_to_event": "Tiempo hasta evento",
    "diagnostic": "Precisión diagnóstica",
    "ordinal": "Ordinal",
    "count": "Conteo/tasa",
    "unknown": "No especificado",
}

DESIGN_OUTCOME_TYPES: dict[str, set[str]] = {
    "mean_estimation": {"continuous"},
    "proportion_estimation": {"binary", "ordinal", "count"},
    "means_independent": {"continuous"},
    "means_paired": {"continuous"},
    "proportions_independent": {"binary"},
    "mcnemar": {"paired_binary", "binary"},
    "superiority_continuous": {"continuous"},
    "superiority_binary": {"binary"},
    "noninferiority_continuous": {"continuous"},
    "noninferiority_binary": {"binary"},
    "equivalence_continuous": {"continuous"},
    "equivalence_binary": {"binary"},
    "odds_ratio": {"binary"},
    "risk_ratio": {"binary"},
    "diagnostic_accuracy": {"diagnostic", "binary"},
    "survival": {"time_to_event"},
}



# -----------------------------------------------------------------------------
# Modelos de datos y utilidades generales
# -----------------------------------------------------------------------------
@dataclass
class SampleSizeResult:
    design_code: str
    design_label: str
    n_raw_total: float
    n_final_total: int
    groups_raw: dict[str, float] = field(default_factory=dict)
    groups_final: dict[str, int] = field(default_factory=dict)
    formula: str = ""
    method: str = ""
    assumptions: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    extra: dict[str, Any] = field(default_factory=dict)
    uses_power: bool = True


DESIGNS: dict[str, str] = {
    "mean_estimation": "Descriptivo · Estimación de una media",
    "proportion_estimation": "Descriptivo · Estimación de una proporción",
    "means_independent": "RCT/Comparativo · Dos medias independientes",
    "means_paired": "RCT/Comparativo · Dos medias apareadas",
    "proportions_independent": "RCT/Comparativo · Dos proporciones independientes",
    "mcnemar": "RCT/Comparativo · Proporciones apareadas (McNemar)",
    "superiority_continuous": "Superioridad · Variable continua",
    "superiority_binary": "Superioridad · Variable dicotómica",
    "noninferiority_continuous": "No inferioridad · Variable continua",
    "noninferiority_binary": "No inferioridad · Variable dicotómica",
    "equivalence_continuous": "Equivalencia · Variable continua (TOST)",
    "equivalence_binary": "Equivalencia · Variable dicotómica (TOST)",
    "odds_ratio": "Asociación · Odds Ratio",
    "risk_ratio": "Asociación · Riesgo Relativo",
    "diagnostic_accuracy": "Validación diagnóstica · Sensibilidad y especificidad",
    "survival": "Supervivencia · Log-rank / Mantel–Cox",
}


ALIASES = {
    "media": "mean_estimation",
    "estimacion_media": "mean_estimation",
    "estimacion_de_media": "mean_estimation",
    "proporcion": "proportion_estimation",
    "estimacion_proporcion": "proportion_estimation",
    "medias_independientes": "means_independent",
    "medias_pareadas": "means_paired",
    "proporciones_independientes": "proportions_independent",
    "mcnemar": "mcnemar",
    "superioridad_continua": "superiority_continuous",
    "superioridad_dicotomica": "superiority_binary",
    "no_inferioridad_continua": "noninferiority_continuous",
    "no_inferioridad_dicotomica": "noninferiority_binary",
    "equivalencia_continua": "equivalence_continuous",
    "equivalencia_dicotomica": "equivalence_binary",
    "odds_ratio": "odds_ratio",
    "or": "odds_ratio",
    "riesgo_relativo": "risk_ratio",
    "rr": "risk_ratio",
    "validacion": "diagnostic_accuracy",
    "validacion_diagnostica": "diagnostic_accuracy",
    "diagnostico": "diagnostic_accuracy",
    "supervivencia": "survival",
    "mantel_cox": "survival",
    "log_rank": "survival",
}


def normalize_token(value: Any) -> str:
    text = str(value or "").strip().lower()
    text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    return re.sub(r"[^a-z0-9]+", "_", text).strip("_")


def clamp_probability(value: float, name: str) -> float:
    value = float(value)
    if not 0 < value < 1:
        raise ValueError(f"{name} debe estar estrictamente entre 0 y 1.")
    return value


def positive(value: float, name: str) -> float:
    value = float(value)
    if value <= 0:
        raise ValueError(f"{name} debe ser mayor que cero.")
    return value


def z_alpha(alpha: float, sided: str) -> float:
    alpha = clamp_probability(alpha, "Error alfa")
    return float(norm.ppf(1 - alpha / 2)) if sided == "Bilateral" else float(norm.ppf(1 - alpha))


def z_beta(power: float) -> float:
    power = clamp_probability(power, "Poder estadístico")
    return float(norm.ppf(power))


def finite_population_correction(n: float, population: int | None) -> float:
    """Corrección exacta para muestreo aleatorio simple sin reemplazo."""
    if not population:
        return n
    if population <= 1:
        raise ValueError("La población finita debe ser mayor que 1.")
    return (population * n) / (population + n - 1)


def inflate_loss(n: float, loss_rate: float) -> int:
    loss_rate = float(loss_rate)
    if not 0 <= loss_rate < 0.95:
        raise ValueError("Las pérdidas esperadas deben estar entre 0% y menos de 95%.")
    return int(math.ceil(n / (1 - loss_rate)))


def continuity_correct_equal_groups(n_uncorrected: float, absolute_difference: float) -> float:
    """Aproximación de Casagrande–Pike–Smith/Fleiss para continuidad."""
    if absolute_difference <= 0:
        raise ValueError("La diferencia entre proporciones debe ser distinta de cero.")
    return (n_uncorrected / 4.0) * (
        1.0 + math.sqrt(1.0 + 2.0 / (n_uncorrected * absolute_difference))
    ) ** 2


def two_proportion_n_per_group(
    p1: float,
    p2: float,
    alpha: float,
    power: float,
    sided: str,
    yates: bool = False,
    null_difference: float = 0.0,
) -> float:
    """Aproximación normal con varianza agrupada bajo H0 y no agrupada bajo H1."""
    p1 = clamp_probability(p1, "Proporción 1")
    p2 = clamp_probability(p2, "Proporción 2")
    za = z_alpha(alpha, sided)
    zb = z_beta(power)
    observed_difference = p1 - p2
    effective_difference = abs(observed_difference - null_difference)
    if effective_difference <= 0:
        raise ValueError("La diferencia efectiva respecto de la hipótesis nula debe ser mayor que cero.")

    # Aproximación estable: promedio de las proporciones esperadas para la varianza nula.
    pbar = (p1 + p2) / 2
    numerator = (
        za * math.sqrt(2 * pbar * (1 - pbar))
        + zb * math.sqrt(p1 * (1 - p1) + p2 * (1 - p2))
    ) ** 2
    n = numerator / (effective_difference**2)
    if yates:
        n = continuity_correct_equal_groups(n, effective_difference)
    return n


def finalize_groups(
    design_code: str,
    design_label: str,
    raw_groups: dict[str, float],
    loss_rate: float,
    formula: str,
    method: str,
    assumptions: list[str],
    warnings: list[str] | None = None,
    extra: dict[str, Any] | None = None,
    uses_power: bool = True,
) -> SampleSizeResult:
    final_groups = {name: inflate_loss(n, loss_rate) for name, n in raw_groups.items()}
    return SampleSizeResult(
        design_code=design_code,
        design_label=design_label,
        n_raw_total=float(sum(raw_groups.values())),
        n_final_total=int(sum(final_groups.values())),
        groups_raw=raw_groups,
        groups_final=final_groups,
        formula=formula,
        method=method,
        assumptions=assumptions,
        warnings=warnings or [],
        extra=extra or {},
        uses_power=uses_power,
    )


# -----------------------------------------------------------------------------
# Motor estadístico
# -----------------------------------------------------------------------------
def calculate_sample_size(params: dict[str, Any], power_override: float | None = None) -> SampleSizeResult:
    code = params["design_code"]
    label = DESIGNS[code]
    alpha = float(params.get("alpha", 0.05))
    power = float(power_override if power_override is not None else params.get("power", 0.80))
    sided = params.get("sided", "Bilateral")
    loss = float(params.get("loss_rate", 0.0))
    za = z_alpha(alpha, sided)
    zb = z_beta(power)
    assumptions = [
        f"α={alpha:.3f}",
        f"contraste {sided.lower()}",
        f"pérdidas esperadas={loss:.1%}",
    ]
    warnings: list[str] = []

    if power < MIN_POWER and code not in {"mean_estimation", "proportion_estimation", "diagnostic_accuracy"}:
        warnings.append("El poder seleccionado es inferior al mínimo convencional de 80%.")

    if code == "mean_estimation":
        sd = positive(params["sd"], "Desviación estándar")
        precision = positive(params["precision"], "Precisión")
        n0 = (za * sd / precision) ** 2
        population = params.get("population") if params.get("finite_population") else None
        n = finite_population_correction(n0, population)
        final = inflate_loss(n, loss)
        assumptions += [f"DE={sd:g}", f"error máximo={precision:g}"]
        if population:
            assumptions.append(f"población finita N={population}")
        return SampleSizeResult(
            code,
            label,
            n,
            final,
            {"Muestra": n},
            {"Muestra": final},
            "n₀=(Z·σ/d)²; nFPC=N·n₀/(N+n₀−1); nfinal=ceil(nFPC/(1−R))",
            "Estimación de una media con precisión absoluta y corrección opcional por población finita.",
            assumptions,
            warnings,
            uses_power=False,
        )

    if code == "proportion_estimation":
        p = clamp_probability(params["p"], "Proporción esperada")
        precision = positive(params["precision"], "Precisión")
        n0 = (za**2) * p * (1 - p) / (precision**2)
        population = params.get("population") if params.get("finite_population") else None
        n = finite_population_correction(n0, population)
        final = inflate_loss(n, loss)
        assumptions += [f"p={p:.3f}", f"error máximo={precision:.3f}"]
        if population:
            assumptions.append(f"población finita N={population}")
        return SampleSizeResult(
            code,
            label,
            n,
            final,
            {"Muestra": n},
            {"Muestra": final},
            "n₀=Z²·p·(1−p)/d²; nFPC=N·n₀/(N+n₀−1); nfinal=ceil(nFPC/(1−R))",
            "Estimación de una proporción con precisión absoluta y corrección opcional por población finita.",
            assumptions,
            warnings,
            uses_power=False,
        )

    if code in {"means_independent", "superiority_continuous"}:
        delta = positive(abs(params["delta"]), "Diferencia clínicamente importante")
        sd = positive(params["sd"], "Desviación estándar")
        n = 2 * (sd**2) * ((za + zb) ** 2) / (delta**2)
        assumptions += [f"DMCI={delta:g}", f"DE común={sd:g}", f"poder={power:.1%}"]
        return finalize_groups(
            code,
            label,
            {"Grupo intervención": n, "Grupo control": n},
            loss,
            "n/grupo=2·σ²·(Zα+Zβ)²/Δ²",
            "Comparación bilateral/unilateral de dos medias independientes con asignación 1:1.",
            assumptions,
            warnings,
        )

    if code == "means_paired":
        delta = positive(abs(params["delta"]), "Diferencia media")
        sd_diff = positive(params["sd_diff"], "DE de las diferencias")
        n = ((za + zb) * sd_diff / delta) ** 2
        final = inflate_loss(n, loss)
        assumptions += [f"diferencia media={delta:g}", f"DE diferencias={sd_diff:g}", f"poder={power:.1%}"]
        return SampleSizeResult(
            code,
            label,
            n,
            final,
            {"Pares/sujetos": n},
            {"Pares/sujetos": final},
            "n=(Zα+Zβ)²·σd²/Δ²",
            "Comparación de una diferencia media dentro de los mismos sujetos.",
            assumptions,
            warnings,
        )

    if code in {"proportions_independent", "superiority_binary"}:
        p1 = params["p1"]
        p2 = params["p2"]
        yates = bool(params.get("yates", False))
        n = two_proportion_n_per_group(p1, p2, alpha, power, sided, yates=yates)
        assumptions += [f"p1={p1:.3f}", f"p2={p2:.3f}", f"poder={power:.1%}"]
        if yates:
            assumptions.append("corrección de continuidad de Yates/Fleiss")
        return finalize_groups(
            code,
            label,
            {"Grupo intervención/expuesto": n, "Grupo control/no expuesto": n},
            loss,
            "n/grupo=[Zα√(2p̄(1−p̄))+Zβ√(p1(1−p1)+p2(1−p2))]²/(p1−p2)²",
            "Comparación de dos proporciones independientes mediante aproximación normal.",
            assumptions,
            warnings,
        )

    if code == "mcnemar":
        p01 = clamp_probability(params["p01"], "Proporción discordante 0→1")
        p10 = clamp_probability(params["p10"], "Proporción discordante 1→0")
        q = p01 + p10
        if q >= 1:
            raise ValueError("La suma de las proporciones discordantes debe ser menor que 1.")
        diff = abs(p01 - p10)
        if diff <= 0:
            raise ValueError("Las proporciones discordantes deben diferir para calcular potencia.")
        n = ((za * math.sqrt(q) + zb * math.sqrt(q - diff**2)) ** 2) / (diff**2)
        if params.get("yates", False):
            n = continuity_correct_equal_groups(n, diff)
            assumptions.append("corrección de continuidad de Yates/Fleiss")
        final = inflate_loss(n, loss)
        assumptions += [f"p01={p01:.3f}", f"p10={p10:.3f}", f"poder={power:.1%}"]
        return SampleSizeResult(
            code,
            label,
            n,
            final,
            {"Pares/sujetos": n},
            {"Pares/sujetos": final},
            "n=[Zα√(p01+p10)+Zβ√(p01+p10−(p01−p10)²)]²/(p01−p10)²",
            "Prueba de McNemar basada en las proporciones discordantes esperadas.",
            assumptions,
            warnings,
        )

    if code == "noninferiority_continuous":
        sd = positive(params["sd"], "Desviación estándar")
        margin = positive(params["margin"], "Margen de no inferioridad")
        expected_diff = float(params.get("expected_diff", 0.0))
        distance = margin + expected_diff  # H0: T-C <= -margin
        if distance <= 0:
            raise ValueError("La diferencia esperada debe estar por encima del límite de no inferioridad (−margen).")
        za1 = float(norm.ppf(1 - alpha))
        n = 2 * sd**2 * (za1 + zb) ** 2 / distance**2
        assumptions += [f"margen={margin:g}", f"diferencia esperada T−C={expected_diff:g}", f"DE={sd:g}"]
        return finalize_groups(
            code,
            label,
            {"Tratamiento": n, "Control": n},
            loss,
            "n/grupo=2·σ²·(Z1−α+Z1−β)²/(Δesperada+M)²",
            "No inferioridad para diferencia de medias, contraste unilateral y asignación 1:1.",
            assumptions,
            warnings,
        )

    if code == "equivalence_continuous":
        sd = positive(params["sd"], "Desviación estándar")
        margin = positive(params["margin"], "Margen de equivalencia")
        expected_diff = abs(float(params.get("expected_diff", 0.0)))
        distance = margin - expected_diff
        if distance <= 0:
            raise ValueError("La diferencia esperada debe quedar estrictamente dentro de ±margen.")
        za1 = float(norm.ppf(1 - alpha))
        n = 2 * sd**2 * (za1 + zb) ** 2 / distance**2
        assumptions += [f"margen simétrico=±{margin:g}", f"|diferencia esperada|={expected_diff:g}", f"DE={sd:g}"]
        return finalize_groups(
            code,
            label,
            {"Tratamiento": n, "Control": n},
            loss,
            "n/grupo=2·σ²·(Z1−α+Z1−β)²/(M−|Δesperada|)²",
            "Equivalencia mediante dos pruebas unilaterales (TOST), asignación 1:1.",
            assumptions,
            warnings,
        )

    if code == "noninferiority_binary":
        p_t = clamp_probability(params["p1"], "Proporción tratamiento")
        p_c = clamp_probability(params["p2"], "Proporción control")
        margin = positive(params["margin"], "Margen de no inferioridad")
        n = two_proportion_n_per_group(
            p_t, p_c, alpha, power, "Unilateral", yates=bool(params.get("yates", False)), null_difference=-margin
        )
        assumptions += [f"pT={p_t:.3f}", f"pC={p_c:.3f}", f"margen absoluto={margin:.3f}"]
        warnings.append("Aproximación normal. En protocolos regulatorios se recomienda confirmar con Farrington–Manning o simulación.")
        return finalize_groups(
            code,
            label,
            {"Tratamiento": n, "Control": n},
            loss,
            "n/grupo≈[Zα√V0+Zβ√V1]²/[(pT−pC)−(−M)]²",
            "No inferioridad de dos proporciones con margen absoluto y contraste unilateral.",
            assumptions,
            warnings,
        )

    if code == "equivalence_binary":
        p_t = clamp_probability(params["p1"], "Proporción tratamiento")
        p_c = clamp_probability(params["p2"], "Proporción control")
        margin = positive(params["margin"], "Margen de equivalencia")
        expected_diff = abs(p_t - p_c)
        distance = margin - expected_diff
        if distance <= 0:
            raise ValueError("La diferencia esperada debe quedar estrictamente dentro de ±margen.")
        pbar = (p_t + p_c) / 2
        za1 = float(norm.ppf(1 - alpha))
        n = (
            za1 * math.sqrt(2 * pbar * (1 - pbar))
            + zb * math.sqrt(p_t * (1 - p_t) + p_c * (1 - p_c))
        ) ** 2 / distance**2
        if params.get("yates", False):
            n = continuity_correct_equal_groups(n, distance)
        assumptions += [f"pT={p_t:.3f}", f"pC={p_c:.3f}", f"margen=±{margin:.3f}"]
        warnings.append("Aproximación TOST normal. En protocolos regulatorios se recomienda confirmar con un método binomial dedicado.")
        return finalize_groups(
            code,
            label,
            {"Tratamiento": n, "Control": n},
            loss,
            "n/grupo≈[Z1−α√V0+Z1−β√V1]²/(M−|pT−pC|)²",
            "Equivalencia de dos proporciones mediante TOST aproximado.",
            assumptions,
            warnings,
        )

    if code == "odds_ratio":
        p_control = clamp_probability(params["p_control"], "Proporción de exposición/evento en controles")
        or_value = positive(params["effect_ratio"], "Odds Ratio")
        if math.isclose(or_value, 1.0, rel_tol=1e-12):
            raise ValueError("El OR debe ser distinto de 1.")
        p_case = (or_value * p_control) / (1 - p_control + or_value * p_control)
        if not 0 < p_case < 1:
            raise ValueError("La combinación de OR y proporción basal produce una probabilidad inválida.")
        n = two_proportion_n_per_group(
            p_case, p_control, alpha, power, sided, yates=bool(params.get("yates", False))
        )
        assumptions += [f"OR={or_value:g}", f"p controles={p_control:.3f}", f"p casos derivada={p_case:.3f}"]
        return finalize_groups(
            code,
            label,
            {"Casos": n, "Controles": n},
            loss,
            "pcasos=(OR·pcontroles)/(1−pcontroles+OR·pcontroles); luego comparación de dos proporciones",
            "Diseño caso-control 1:1 basado en OR mínimo detectable.",
            assumptions,
            warnings,
            extra={"derived_probability": p_case},
        )

    if code == "risk_ratio":
        p_control = clamp_probability(params["p_control"], "Riesgo basal en no expuestos/control")
        rr_value = positive(params["effect_ratio"], "Riesgo Relativo")
        if math.isclose(rr_value, 1.0, rel_tol=1e-12):
            raise ValueError("El RR debe ser distinto de 1.")
        p_exposed = rr_value * p_control
        if not 0 < p_exposed < 1:
            raise ValueError("RR × riesgo basal debe quedar entre 0 y 1.")
        n = two_proportion_n_per_group(
            p_exposed, p_control, alpha, power, sided, yates=bool(params.get("yates", False))
        )
        assumptions += [f"RR={rr_value:g}", f"riesgo basal={p_control:.3f}", f"riesgo expuesto derivado={p_exposed:.3f}"]
        return finalize_groups(
            code,
            label,
            {"Expuestos/intervención": n, "No expuestos/control": n},
            loss,
            "pexpuestos=RR·pcontrol; luego comparación de dos proporciones",
            "Cohorte o ensayo 1:1 basado en RR mínimo detectable.",
            assumptions,
            warnings,
            extra={"derived_probability": p_exposed},
        )

    if code == "diagnostic_accuracy":
        sensitivity = clamp_probability(params["sensitivity"], "Sensibilidad esperada")
        specificity = clamp_probability(params["specificity"], "Especificidad esperada")
        prevalence = clamp_probability(params["prevalence"], "Prevalencia")
        d_se = positive(params["precision_se"], "Precisión para sensibilidad")
        d_sp = positive(params["precision_sp"], "Precisión para especificidad")
        z = float(norm.ppf(1 - alpha / 2))
        n_se_total = z**2 * sensitivity * (1 - sensitivity) / (d_se**2 * prevalence)
        n_sp_total = z**2 * specificity * (1 - specificity) / (d_sp**2 * (1 - prevalence))
        selected_raw = max(n_se_total, n_sp_total)
        selected_final = inflate_loss(selected_raw, loss)
        se_final = inflate_loss(n_se_total, loss)
        sp_final = inflate_loss(n_sp_total, loss)
        assumptions += [
            f"Se={sensitivity:.3f}",
            f"Sp={specificity:.3f}",
            f"prevalencia={prevalence:.3f}",
            f"precisión Se=±{d_se:.3f}",
            f"precisión Sp=±{d_sp:.3f}",
        ]
        warnings.append("Este método estima precisión de intervalos de confianza; el parámetro de poder no interviene.")
        return SampleSizeResult(
            code,
            label,
            selected_raw,
            selected_final,
            {"Requerimiento por sensibilidad": n_se_total, "Requerimiento por especificidad": n_sp_total},
            {"Sensibilidad": se_final, "Especificidad": sp_final, "Recomendado (mayor)": selected_final},
            "NSe=Z²·Se·(1−Se)/(dSe²·Prev); NSp=Z²·Sp·(1−Sp)/(dSp²·(1−Prev)); N=max(NSe,NSp)",
            "Método de precisión diagnóstica ajustado por prevalencia (enfoque de Buderer).",
            assumptions,
            warnings,
            extra={"n_sensitivity": se_final, "n_specificity": sp_final},
            uses_power=False,
        )

    if code == "survival":
        hr = positive(params["hazard_ratio"], "Hazard Ratio")
        if math.isclose(hr, 1.0, rel_tol=1e-12):
            raise ValueError("El HR debe ser distinto de 1.")
        mortality_control = clamp_probability(params["mortality_control"], "Mortalidad acumulada control")
        mortality_treatment = clamp_probability(params["mortality_treatment"], "Mortalidad acumulada tratamiento")
        allocation_t = clamp_probability(params.get("allocation_treatment", 0.5), "Fracción de asignación al tratamiento")
        event_probability = allocation_t * mortality_treatment + (1 - allocation_t) * mortality_control
        events = (za + zb) ** 2 / ((math.log(hr) ** 2) * allocation_t * (1 - allocation_t))
        total_raw = events / event_probability
        n_t = total_raw * allocation_t
        n_c = total_raw * (1 - allocation_t)
        assumptions += [
            f"HR={hr:g}",
            f"mortalidad control={mortality_control:.1%}",
            f"mortalidad tratamiento={mortality_treatment:.1%}",
            f"fracción tratamiento={allocation_t:.2f}",
            f"eventos requeridos≈{math.ceil(events)}",
        ]
        warnings.append("La conversión de eventos a sujetos supone mortalidad acumulada promedio y riesgos proporcionales.")
        return finalize_groups(
            code,
            label,
            {"Tratamiento": n_t, "Control": n_c},
            loss,
            "D=(Zα+Zβ)²/[q(1−q)·ln(HR)²]; N=D/probabilidad promedio de evento",
            "Número de eventos de Schoenfeld para log-rank/Mantel–Cox y conversión aproximada a participantes.",
            assumptions,
            warnings,
            extra={"required_events": int(math.ceil(events)), "event_probability": event_probability},
        )

    raise ValueError("Diseño no implementado.")



def _entry_number(entry: dict[str, Any], keys: tuple[str, ...], default: float | None = None) -> float | None:
    for key in keys:
        if entry.get(key) not in (None, ""):
            return _parse_locale_number(entry[key])
    return default


def _as_probability(value: float, name: str) -> float:
    numeric = float(value)
    if 1 < numeric <= 100:
        numeric /= 100
    return clamp_probability(numeric, name)


def build_params_for_structured_outcome(
    entry: dict[str, Any],
    base_params: dict[str, Any],
) -> dict[str, Any]:
    outcome_type = infer_outcome_type(entry.get("tipo_outcome", ""))
    design_value = entry.get("tipo_diseno")
    if design_value:
        code = resolve_design(design_value)
    else:
        code = suggested_design_for_outcome(outcome_type, base_params.get("design_code")) or base_params["design_code"]

    alpha = _entry_number(entry, ("error_alfa",), float(base_params.get("alpha", 0.05)))
    power = _entry_number(entry, ("poder_estadistico",), float(base_params.get("power", 0.80)))
    loss = _entry_number(entry, ("perdidas_esperadas",), float(base_params.get("loss_rate", 0.0)))
    assert alpha is not None and power is not None and loss is not None
    if alpha > 1:
        alpha /= 100
    if power > 1:
        power /= 100
    if loss > 1:
        loss /= 100

    params: dict[str, Any] = {
        "design_code": code,
        "alpha": alpha,
        "power": power,
        "sided": entry.get("sided", base_params.get("sided", "Bilateral")),
        "loss_rate": loss,
        "outcome": {
            "outcome_primario": entry.get("outcome_primario", ""),
            "definicion_outcome": entry.get("definicion_outcome", ""),
            "tipo_outcome": outcome_type,
            "tipo_outcome_label": OUTCOME_TYPE_LABELS.get(outcome_type, outcome_type),
            "unidad_outcome": entry.get("unidad_outcome", ""),
            "momento_evaluacion": entry.get("momento_evaluacion", ""),
            "columna_dataset_outcome": entry.get("columna_dataset_outcome", ""),
            "valor_evento": entry.get("valor_evento", ""),
            "confirmado": True,
        },
    }
    effect = _entry_number(entry, ("efecto_esperado", "delta", "margen"))
    variability = _entry_number(entry, ("variabilidad_estimada", "sd", "desviacion_estandar"))

    if code == "mean_estimation":
        params["sd"] = positive(variability if variability is not None else 0, "Variabilidad estimada")
        params["precision"] = positive(
            _entry_number(entry, ("precision",), effect) or 0,
            "Precisión",
        )
        params["finite_population"] = bool(entry.get("finite_population", False))
        if params["finite_population"]:
            params["population"] = int(_entry_number(entry, ("poblacion",), 0) or 0)
    elif code == "proportion_estimation":
        proportion = _entry_number(entry, ("proporcion_esperada", "p"), effect)
        precision = _entry_number(entry, ("precision",), variability)
        params["p"] = _as_probability(proportion or 0, "Proporción esperada")
        params["precision"] = positive(precision or 0, "Precisión")
        if params["precision"] > 1:
            params["precision"] /= 100
        params["finite_population"] = bool(entry.get("finite_population", False))
        if params["finite_population"]:
            params["population"] = int(_entry_number(entry, ("poblacion",), 0) or 0)
    elif code in {"means_independent", "superiority_continuous"}:
        params["delta"] = positive(effect or 0, "Efecto esperado")
        params["sd"] = positive(variability or 0, "Variabilidad estimada")
    elif code == "means_paired":
        params["delta"] = positive(effect or 0, "Efecto esperado")
        params["sd_diff"] = positive(variability or 0, "Variabilidad de las diferencias")
    elif code in {"proportions_independent", "superiority_binary"}:
        p2 = _entry_number(entry, ("proporcion_grupo_2", "proporcion_control", "p2"), variability)
        p1 = _entry_number(entry, ("proporcion_grupo_1", "p1"))
        p2_prob = _as_probability(p2 or 0, "Proporción control")
        if p1 is None:
            diff = float(effect or 0)
            if abs(diff) > 1:
                diff /= 100
            p1 = p2_prob + diff
        params["p1"] = _as_probability(p1, "Proporción grupo 1")
        params["p2"] = p2_prob
        params["yates"] = bool(entry.get("yates", base_params.get("yates", False)))
    elif code == "mcnemar":
        p01 = _entry_number(entry, ("p01",))
        p10 = _entry_number(entry, ("p10",))
        if p01 is None or p10 is None:
            discordance = float(variability or 0)
            difference = float(effect or 0)
            if discordance > 1:
                discordance /= 100
            if difference > 1:
                difference /= 100
            p01 = (discordance + difference) / 2
            p10 = (discordance - difference) / 2
        params["p01"] = _as_probability(p01, "p01")
        params["p10"] = _as_probability(p10, "p10")
        params["yates"] = bool(entry.get("yates", base_params.get("yates", False)))
    elif code in {"noninferiority_continuous", "equivalence_continuous"}:
        params["sd"] = positive(variability or 0, "Variabilidad estimada")
        params["margin"] = positive(_entry_number(entry, ("margen",), effect) or 0, "Margen")
        params["expected_diff"] = float(_entry_number(entry, ("expected_diff", "diferencia_esperada"), 0.0) or 0.0)
        params["sided"] = "Unilateral"
    elif code in {"noninferiority_binary", "equivalence_binary"}:
        baseline = _entry_number(entry, ("proporcion_grupo_2", "proporcion_control", "p2"), variability)
        p2_prob = _as_probability(baseline or 0, "Proporción control")
        p1_value = _entry_number(entry, ("proporcion_grupo_1", "p1"), p2_prob)
        margin = _entry_number(entry, ("margen",), effect)
        params["p1"] = _as_probability(p1_value or 0, "Proporción tratamiento")
        params["p2"] = p2_prob
        params["margin"] = positive(margin or 0, "Margen")
        if params["margin"] > 1:
            params["margin"] /= 100
        params["yates"] = bool(entry.get("yates", base_params.get("yates", False)))
        params["sided"] = "Unilateral"
    elif code in {"odds_ratio", "risk_ratio"}:
        baseline = _entry_number(entry, ("proporcion_control", "p_control"), variability)
        ratio_key = "odds_ratio" if code == "odds_ratio" else "riesgo_relativo"
        ratio = _entry_number(entry, (ratio_key, "effect_ratio"), effect)
        params["p_control"] = _as_probability(baseline or 0, "Riesgo basal")
        params["effect_ratio"] = positive(ratio or 0, "Medida de asociación")
        params["yates"] = bool(entry.get("yates", base_params.get("yates", False)))
    elif code == "diagnostic_accuracy":
        prevalence = _entry_number(entry, ("prevalencia",), variability)
        sensitivity = _entry_number(entry, ("sensibilidad",), effect)
        specificity = _entry_number(entry, ("especificidad",), effect)
        params["prevalence"] = _as_probability(prevalence or 0, "Prevalencia")
        params["sensitivity"] = _as_probability(sensitivity or 0, "Sensibilidad")
        params["specificity"] = _as_probability(specificity or 0, "Especificidad")
        params["precision_se"] = _entry_number(entry, ("precision_sensibilidad",), base_params.get("precision_se", 0.05))
        params["precision_sp"] = _entry_number(entry, ("precision_especificidad",), base_params.get("precision_sp", 0.05))
        if params["precision_se"] > 1:
            params["precision_se"] /= 100
        if params["precision_sp"] > 1:
            params["precision_sp"] /= 100
    elif code == "survival":
        hr = _entry_number(entry, ("hazard_ratio",), effect)
        mortality_control = _entry_number(entry, ("mortalidad_control",), variability)
        params["hazard_ratio"] = positive(hr or 0, "Hazard ratio")
        params["mortality_control"] = _as_probability(mortality_control or 0, "Mortalidad control")
        treatment = _entry_number(entry, ("mortalidad_tratamiento",))
        if treatment is None:
            treatment = min(max(params["mortality_control"] * params["hazard_ratio"], 0.001), 0.999)
        params["mortality_treatment"] = _as_probability(treatment, "Mortalidad tratamiento")
        params["allocation_treatment"] = float(
            _entry_number(entry, ("allocation_treatment", "asignacion_tratamiento"), base_params.get("allocation_treatment", 0.50))
            or 0.50
        )
    else:
        raise ValueError(f"No se pudo construir el escenario para {code}.")

    return params


def calculate_structured_outcome_scenarios(
    outcomes: list[dict[str, Any]],
    base_params: dict[str, Any],
) -> tuple[pd.DataFrame, list[str]]:
    rows: list[dict[str, Any]] = []
    warnings: list[str] = []
    for entry in outcomes:
        if not (entry.get("es_primario") or entry.get("es_coprimario")):
            continue
        name = str(entry.get("outcome_primario", "Outcome sin nombre"))
        try:
            scenario_params = build_params_for_structured_outcome(entry, base_params)
            scenario_result = calculate_sample_size(scenario_params)
            rows.append(
                {
                    "Outcome": name,
                    "Rol": "Coprimario" if entry.get("es_coprimario") else "Primario",
                    "Tipo": scenario_params["outcome"].get("tipo_outcome_label", ""),
                    "Diseño": DESIGNS[scenario_params["design_code"]],
                    "N total": scenario_result.n_final_total,
                    "Estado": "Calculado",
                }
            )
        except Exception as exc:
            warnings.append(f"{name}: {exc}")
            rows.append(
                {
                    "Outcome": name,
                    "Rol": "Coprimario" if entry.get("es_coprimario") else "Primario",
                    "Tipo": OUTCOME_TYPE_LABELS.get(infer_outcome_type(entry.get("tipo_outcome", "")), "No especificado"),
                    "Diseño": str(entry.get("tipo_diseno", "No especificado")),
                    "N total": "—",
                    "Estado": "Parámetros insuficientes",
                }
            )
    return pd.DataFrame(rows), warnings


# -----------------------------------------------------------------------------
# Importación de datos y protocolos
# -----------------------------------------------------------------------------
PROTOCOL_FIELD_ALIASES: dict[str, tuple[str, ...]] = {
    "tipo_diseno": (
        "tipo de diseño",
        "tipo de diseno",
        "diseño del estudio",
        "diseno del estudio",
        "diseño estadístico",
        "diseno estadistico",
        "diseño metodológico",
        "diseno metodologico",
    ),
    "error_alfa": (
        "error alfa",
        "nivel alfa",
        "nivel de significación",
        "nivel de significacion",
        "significancia alfa",
        "alpha",
        "alfa",
    ),
    "poder_estadistico": (
        "poder estadístico",
        "poder estadistico",
        "potencia estadística",
        "potencia estadistica",
        "potencia del estudio",
        "poder del estudio",
        "1-beta",
        "1 - beta",
    ),
    "efecto_esperado": (
        "efecto esperado",
        "tamaño del efecto",
        "tamano del efecto",
        "diferencia mínima clínicamente importante",
        "diferencia minima clinicamente importante",
        "diferencia clínicamente importante",
        "diferencia clinicamente importante",
        "diferencia esperada",
        "dmci",
        "delta esperado",
    ),
    "variabilidad_estimada": (
        "variabilidad estimada",
        "desviación estándar estimada",
        "desviacion estandar estimada",
        "desvío estándar estimado",
        "desvio estandar estimado",
        "desviación estándar",
        "desviacion estandar",
        "desvío estándar",
        "desvio estandar",
        "sigma estimada",
        "sigma",
    ),
    "perdidas_esperadas": (
        "pérdidas esperadas",
        "perdidas esperadas",
        "pérdidas de seguimiento",
        "perdidas de seguimiento",
        "tasa de pérdidas",
        "tasa de perdidas",
        "abandono esperado",
    ),
    "precision": (
        "precisión absoluta",
        "precision absoluta",
        "error máximo aceptable",
        "error maximo aceptable",
        "precisión esperada",
        "precision esperada",
    ),
    "margen": (
        "margen de no inferioridad",
        "margen de equivalencia",
        "margen clínico",
        "margen clinico",
        "margen absoluto",
    ),
    "proporcion_grupo_1": (
        "proporción grupo 1",
        "proporcion grupo 1",
        "proporción tratamiento",
        "proporcion tratamiento",
        "tasa tratamiento",
    ),
    "proporcion_grupo_2": (
        "proporción grupo 2",
        "proporcion grupo 2",
        "proporción control",
        "proporcion control",
        "tasa control",
    ),
    "proporcion_esperada": ("proporción esperada", "proporcion esperada"),
    "proporcion_control": (
        "riesgo basal en control",
        "proporción basal en control",
        "proporcion basal en control",
        "proporción de expuestos en controles",
        "proporcion de expuestos en controles",
    ),
    "prevalencia": ("prevalencia", "prevalencia de la enfermedad"),
    "sensibilidad": ("sensibilidad esperada", "sensibilidad"),
    "especificidad": ("especificidad esperada", "especificidad"),
    "precision_sensibilidad": (
        "precisión para sensibilidad",
        "precision para sensibilidad",
        "error de sensibilidad",
    ),
    "precision_especificidad": (
        "precisión para especificidad",
        "precision para especificidad",
        "error de especificidad",
    ),
    "odds_ratio": ("odds ratio", "razón de momios", "razon de momios"),
    "riesgo_relativo": ("riesgo relativo", "relative risk"),
    "hazard_ratio": ("hazard ratio", "razón de riesgos", "razon de riesgos"),
    "mortalidad_control": (
        "mortalidad acumulada control",
        "mortalidad en control",
        "tasa de mortalidad control",
    ),
    "mortalidad_tratamiento": (
        "mortalidad acumulada tratamiento",
        "mortalidad en tratamiento",
        "tasa de mortalidad tratamiento",
    ),
    "poblacion": ("tamaño de la población", "tamano de la poblacion", "población finita", "poblacion finita"),
    "p01": ("discordantes 0 a 1", "discordantes 0→1", "p01"),
    "p10": ("discordantes 1 a 0", "discordantes 1→0", "p10"),
    "outcome_primario": (
        "desenlace primario",
        "desenlace principal",
        "variable de resultado primaria",
        "variable de resultado principal",
        "variable primaria",
        "variable principal",
        "outcome primario",
        "primary outcome",
        "endpoint primario",
        "primary endpoint",
        "criterio principal de valoración",
        "criterio principal de valoracion",
        "resultado primario",
    ),
    "definicion_outcome": (
        "definición del desenlace",
        "definicion del desenlace",
        "definición del outcome",
        "definicion del outcome",
        "definición operacional",
        "definicion operacional",
        "definición de la variable primaria",
        "definicion de la variable primaria",
    ),
    "tipo_outcome": (
        "tipo de outcome",
        "tipo de desenlace",
        "naturaleza del desenlace",
        "escala del desenlace",
        "tipo de variable primaria",
    ),
    "unidad_outcome": (
        "unidad del outcome",
        "unidad del desenlace",
        "unidad de medida",
        "unidades del resultado",
    ),
    "momento_evaluacion": (
        "momento de evaluación",
        "momento de evaluacion",
        "tiempo de evaluación",
        "tiempo de evaluacion",
        "punto temporal",
        "timepoint",
        "follow-up del outcome",
    ),
    "columna_dataset_outcome": (
        "columna del outcome",
        "columna del dataset",
        "variable en la base",
        "nombre de columna",
        "campo del dataset",
    ),
    "valor_evento": (
        "valor del evento",
        "valor positivo",
        "evento codificado como",
        "categoría positiva",
        "categoria positiva",
    ),
    "outcome_secundarios": (
        "desenlaces secundarios",
        "outcomes secundarios",
        "endpoints secundarios",
        "variables secundarias",
        "resultados secundarios",
    ),
}

TEXT_PROTOCOL_FIELDS = {
    "tipo_diseno",
    "outcome_primario",
    "definicion_outcome",
    "tipo_outcome",
    "unidad_outcome",
    "momento_evaluacion",
    "columna_dataset_outcome",
    "valor_evento",
    "outcome_secundarios",
}
NUMERIC_PROTOCOL_FIELDS = set(PROTOCOL_FIELD_ALIASES) - TEXT_PROTOCOL_FIELDS


def _accentless_text(value: Any) -> str:
    text = str(value or "").replace("\u00a0", " ")
    text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    return text.lower()


def _parse_locale_number(value: Any) -> float:
    if isinstance(value, (int, float, np.integer, np.floating)):
        number = float(value)
        if not math.isfinite(number):
            raise ValueError("El valor numérico no es finito.")
        return number

    text = str(value or "").strip().replace("\u00a0", " ")
    # Evita interpretar la notación de potencia (1−β) como el valor numérico.
    text = re.sub(r"^\s*\([^)]*\)\s*", "", text)
    text = re.sub(r"^\s*1\s*[-−–]\s*(?:beta|β)\s*[:=]?\s*", "", text, flags=re.IGNORECASE)
    match = re.search(r"[-+]?\d(?:[\d\s.,]*\d)?|[-+]?\d", text)
    if not match:
        raise ValueError(f"No se encontró un número en {value!r}.")
    token = match.group(0).replace(" ", "")
    if "," in token and "." in token:
        decimal_sep = "," if token.rfind(",") > token.rfind(".") else "."
        thousands_sep = "." if decimal_sep == "," else ","
        token = token.replace(thousands_sep, "").replace(decimal_sep, ".")
    elif "," in token:
        parts = token.split(",")
        token = "".join(parts[:-1]) + "." + parts[-1] if len(parts[-1]) <= 4 else "".join(parts)
    elif token.count(".") > 1:
        parts = token.split(".")
        token = "".join(parts[:-1]) + "." + parts[-1]
    number = float(token)
    if not math.isfinite(number):
        raise ValueError("El valor numérico no es finito.")
    return number


def _extract_docx_text(raw: bytes) -> str:
    document = Document(io.BytesIO(raw))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            blocks.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append("\t".join(cells))
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                blocks.append(paragraph.text.strip())
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                blocks.append(paragraph.text.strip())
    return "\n".join(blocks)


def _extract_pdf_text(raw: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ValueError("El PDF está cifrado y no pudo abrirse sin contraseña.") from exc
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join(page for page in pages if page)


def _extract_legacy_doc_text(raw: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(raw)
        temp_path = tmp.name
    try:
        for command_name in ("antiword", "catdoc"):
            executable = shutil.which(command_name)
            if executable:
                completed = subprocess.run(
                    [executable, temp_path],
                    check=False,
                    capture_output=True,
                    timeout=30,
                )
                if completed.returncode == 0 and completed.stdout.strip():
                    for encoding in ("utf-8", "latin-1"):
                        try:
                            text = completed.stdout.decode(encoding)
                            if text.strip():
                                return text, warnings
                        except UnicodeDecodeError:
                            continue

        # Rescate heurístico para Word binario cuando antiword/catdoc no están instalados.
        ascii_chunks = [m.decode("latin-1", errors="ignore") for m in re.findall(rb"[\x20-\x7e\xa0-\xff]{5,}", raw)]
        utf16_chunks: list[str] = []
        for match in re.findall(rb"(?:[\x20-\x7e\xa0-\xff]\x00){5,}", raw):
            try:
                utf16_chunks.append(match.decode("utf-16le", errors="ignore"))
            except Exception:
                pass
        text = "\n".join(dict.fromkeys(chunk.strip() for chunk in utf16_chunks + ascii_chunks if chunk.strip()))
        warnings.append(
            "El archivo .DOC legado se leyó mediante recuperación heurística. Para máxima confiabilidad, guárdelo como DOCX."
        )
        return text, warnings
    finally:
        try:
            os.unlink(temp_path)
        except OSError:
            pass


def extract_protocol_text(raw: bytes, suffix: str) -> tuple[str, list[str]]:
    warnings: list[str] = []
    if suffix == ".pdf":
        text = _extract_pdf_text(raw)
        if len(text.strip()) < 30:
            raise ValueError(
                "No se detectó texto utilizable en el PDF. Si es un documento escaneado, aplique OCR o conviértalo a PDF con texto seleccionable."
            )
        return text, warnings
    if suffix in {".docx", ".docm"}:
        text = _extract_docx_text(raw)
        if len(text.strip()) < 20:
            raise ValueError("El archivo Word no contiene texto utilizable.")
        return text, warnings
    if suffix == ".doc":
        text, legacy_warnings = _extract_legacy_doc_text(raw)
        warnings.extend(legacy_warnings)
        if len(text.strip()) < 20:
            raise ValueError("No fue posible recuperar texto del archivo .DOC. Conviértalo a DOCX.")
        return text, warnings
    raise ValueError("Formato documental no compatible.")


def _value_after_alias(lines: list[str], aliases: tuple[str, ...]) -> str | None:
    normalized_aliases = sorted({_accentless_text(alias).strip() for alias in aliases}, key=len, reverse=True)
    for index, line in enumerate(lines):
        normalized_line = _accentless_text(line)
        for alias in normalized_aliases:
            position = normalized_line.find(alias)
            if position < 0:
                continue
            tail = line[position + len(alias):].strip()
            tail = re.sub(r"^[\s\t:=\-–—·•()\[\]]+", "", tail).strip()
            if tail:
                return tail
            if index + 1 < len(lines):
                next_line = lines[index + 1].strip()
                if next_line:
                    return next_line
    return None


def infer_design_from_text(text: str) -> str | None:
    token = normalize_token(text)

    def has(*parts: str) -> bool:
        return any(normalize_token(part) in token for part in parts)

    binary_hint = has("dicotomica", "binaria", "proporcion", "evento", "respuesta si no", "tasa")
    continuous_hint = has("continua", "media", "promedio", "desviacion estandar")

    if has("mcnemar"):
        return "mcnemar"
    if has("validacion diagnostica", "precision diagnostica") or (has("sensibilidad") and has("especificidad")):
        return "diagnostic_accuracy"
    if has("mantel cox", "log rank", "logrank", "curvas de supervivencia", "hazard ratio"):
        return "survival"
    if has("odds ratio", "razon de momios"):
        return "odds_ratio"
    if has("riesgo relativo", "relative risk"):
        return "risk_ratio"
    if has("no inferioridad", "no-inferioridad"):
        return "noninferiority_binary" if binary_hint and not continuous_hint else "noninferiority_continuous"
    if has("equivalencia", "tost"):
        return "equivalence_binary" if binary_hint and not continuous_hint else "equivalence_continuous"
    if has("superioridad"):
        return "superiority_binary" if binary_hint and not continuous_hint else "superiority_continuous"
    if has("medias apareadas", "media apareada", "antes y despues", "medidas pareadas"):
        return "means_paired"
    if has("proporciones apareadas", "datos apareados dicotomicos"):
        return "mcnemar"
    if has("medias independientes", "dos medias independientes", "comparacion de medias independientes"):
        return "means_independent"
    if has("proporciones independientes", "dos proporciones independientes", "comparacion de proporciones"):
        return "proportions_independent"
    if has("estimacion de una proporcion", "estimacion de proporcion", "prevalencia con precision"):
        return "proportion_estimation"
    if has("estimacion de una media", "estimacion de media", "media con precision"):
        return "mean_estimation"
    return None



def _clean_outcome_value(value: Any) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip(" \t:;-–—.\n")
    text = re.sub(
        r"^(?:sera|será|es|fue|se define como|se definio como|se definió como|consiste en|consistira en|consistirá en)\s+",
        "",
        text,
        flags=re.IGNORECASE,
    )
    return text.strip(" .;:")


def infer_outcome_type(value: Any, design_code: str | None = None) -> str:
    token = normalize_token(value)

    explicit_patterns = {
        "time_to_event": ("tiempo_hasta_evento", "time_to_event", "supervivencia", "hazard", "mortalidad_hasta"),
        "diagnostic": ("diagnostica", "diagnostico", "sensibilidad", "especificidad", "area_bajo_la_curva", "auc"),
        "paired_binary": ("binaria_apareada", "dicotomica_apareada", "mcnemar"),
        "binary": ("binaria", "binario", "dicotomica", "dicotomico", "si_no", "presencia_ausencia", "incidencia", "evento", "mortalidad"),
        "ordinal": ("ordinal", "likert", "grado", "clase_funcional"),
        "count": ("conteo", "numero_de_eventos", "tasa_de_eventos", "recuento"),
        "continuous": ("continua", "continuo", "cuantitativa", "cuantitativo", "media", "promedio", "cambio", "reduccion", "incremento"),
    }
    for outcome_type, patterns in explicit_patterns.items():
        if any(pattern in token for pattern in patterns):
            return outcome_type

    if design_code:
        expected = DESIGN_OUTCOME_TYPES.get(design_code, set())
        if len(expected) == 1:
            return next(iter(expected))
        if design_code == "mcnemar":
            return "paired_binary"
        if design_code == "diagnostic_accuracy":
            return "diagnostic"

    return "unknown"


def infer_outcome_unit(text: Any) -> str:
    value = str(text or "")
    patterns = [
        r"\bmm\s*hg\b",
        r"\bmg\s*/\s*d[lL]\b",
        r"\bmmol\s*/\s*[lL]\b",
        r"\bg\s*/\s*d[lL]\b",
        r"\bml\s*/\s*min(?:\s*/\s*1[,.]73\s*m(?:2|²))?\b",
        r"\bkg\s*/\s*m(?:2|²)\b",
        r"\bm/s\b",
        r"\bms\b",
        r"\bbpm\b",
        r"\bkg\b",
        r"\bcm\b",
        r"\bmm\b",
        r"\b%\b",
        r"\bpuntos?\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, value, flags=re.IGNORECASE)
        if match:
            return re.sub(r"\s+", "", match.group(0)).replace("²", "2")
    return ""


def infer_outcome_timepoint(text: Any) -> str:
    value = re.sub(r"\s+", " ", str(text or ""))
    patterns = [
        r"(?:a|al|a las|en la|en el)\s+(?:semana|día|dia|mes|año|ano)\s*\d+",
        r"(?:a|al|a las|en)\s+\d+(?:[.,]\d+)?\s*(?:horas?|días?|dias?|semanas?|meses?|años?|anos?)",
        r"(?:tras|después de|despues de)\s+\d+(?:[.,]\d+)?\s*(?:horas?|días?|dias?|semanas?|meses?|años?|anos?)",
        r"(?:hasta|durante)\s+\d+(?:[.,]\d+)?\s*(?:horas?|días?|dias?|semanas?|meses?|años?|anos?)",
        r"(?:basal|fin del seguimiento|alta hospitalaria)",
    ]
    for pattern in patterns:
        match = re.search(pattern, value, flags=re.IGNORECASE)
        if match:
            return match.group(0).strip(" .;:")
    return ""


def design_compatible_with_outcome(design_code: str, outcome_type: str) -> bool:
    if outcome_type == "unknown":
        return True
    return outcome_type in DESIGN_OUTCOME_TYPES.get(design_code, {outcome_type})


def suggested_design_for_outcome(outcome_type: str, current_design: str | None = None) -> str | None:
    if current_design and design_compatible_with_outcome(current_design, outcome_type):
        return current_design
    suggestions = {
        "continuous": "means_independent",
        "binary": "proportions_independent",
        "paired_binary": "mcnemar",
        "time_to_event": "survival",
        "diagnostic": "diagnostic_accuracy",
        "ordinal": "proportion_estimation",
        "count": "proportion_estimation",
    }
    return suggestions.get(outcome_type)


def _canonicalize_outcome_entry(entry: dict[str, Any]) -> dict[str, Any]:
    aliases = {
        "nombre": "outcome_primario",
        "name": "outcome_primario",
        "outcome": "outcome_primario",
        "desenlace": "outcome_primario",
        "endpoint": "outcome_primario",
        "definicion": "definicion_outcome",
        "definition": "definicion_outcome",
        "tipo": "tipo_outcome",
        "outcome_type": "tipo_outcome",
        "unidad": "unidad_outcome",
        "unit": "unidad_outcome",
        "momento": "momento_evaluacion",
        "timepoint": "momento_evaluacion",
        "columna": "columna_dataset_outcome",
        "column": "columna_dataset_outcome",
        "valor_positivo": "valor_evento",
        "positive_value": "valor_evento",
        "rol": "rol_outcome",
        "role": "rol_outcome",
        "primary": "es_primario",
        "coprimary": "es_coprimario",
    }
    canonical: dict[str, Any] = {}
    for key, value in entry.items():
        token = normalize_token(key)
        canonical[aliases.get(token, token)] = value
    if canonical.get("outcome_primario"):
        canonical["outcome_primario"] = _clean_outcome_value(canonical["outcome_primario"])
    role_token = normalize_token(canonical.get("rol_outcome", ""))
    canonical["es_primario"] = bool(canonical.get("es_primario")) or role_token in {"primario", "primary"}
    canonical["es_coprimario"] = bool(canonical.get("es_coprimario")) or role_token in {
        "coprimario",
        "co_primario",
        "coprimary",
    }
    return canonical


def extract_structured_outcomes(data: dict[str, Any]) -> list[dict[str, Any]]:
    for key in ("outcomes", "desenlaces", "endpoints", "variables_resultado"):
        raw = data.get(key)
        if isinstance(raw, list):
            return [_canonicalize_outcome_entry(item) for item in raw if isinstance(item, dict)]
    return []


def choose_primary_outcome(outcomes: list[dict[str, Any]]) -> dict[str, Any] | None:
    if not outcomes:
        return None
    for item in outcomes:
        if item.get("es_primario"):
            return item
    for item in outcomes:
        if not item.get("es_coprimario"):
            return item
    return outcomes[0]


def enrich_outcome_fields(protocol: dict[str, Any], source_text: str = "") -> dict[str, Any]:
    outcome_name = _clean_outcome_value(protocol.get("outcome_primario", ""))
    if outcome_name:
        protocol["outcome_primario"] = outcome_name
    combined = " ".join(
        str(protocol.get(key, ""))
        for key in ("outcome_primario", "definicion_outcome", "tipo_outcome", "unidad_outcome", "momento_evaluacion")
    )
    if not protocol.get("definicion_outcome") and outcome_name:
        protocol["definicion_outcome"] = outcome_name
    design_code = None
    if protocol.get("tipo_diseno"):
        try:
            design_code = resolve_design(protocol["tipo_diseno"])
        except Exception:
            design_code = infer_design_from_text(str(protocol["tipo_diseno"]))
    if not protocol.get("tipo_outcome"):
        protocol["tipo_outcome"] = infer_outcome_type(combined or source_text, design_code)
    else:
        protocol["tipo_outcome"] = infer_outcome_type(protocol["tipo_outcome"], design_code)
    if not protocol.get("unidad_outcome"):
        unit = infer_outcome_unit(combined or source_text)
        if unit:
            protocol["unidad_outcome"] = unit
    if not protocol.get("momento_evaluacion"):
        timepoint = infer_outcome_timepoint(combined or source_text)
        if timepoint:
            protocol["momento_evaluacion"] = timepoint
    return protocol


def _is_binary_series(series: pd.Series) -> bool:
    values = series.dropna()
    if values.empty:
        return False
    unique_count = int(values.nunique(dropna=True))
    return unique_count <= 2 or pd.api.types.is_bool_dtype(values)


def _is_low_cardinality_series(series: pd.Series, max_categories: int = 10) -> bool:
    values = series.dropna()
    return not values.empty and int(values.nunique(dropna=True)) <= max_categories


def outcome_column_candidates(
    df: pd.DataFrame,
    outcome_name: str,
    outcome_type: str,
    preferred_column: str | None = None,
) -> list[dict[str, Any]]:
    target = normalize_token(outcome_name)
    target_tokens = {token for token in target.split("_") if len(token) > 1}
    candidates: list[dict[str, Any]] = []
    for column in df.columns:
        column_text = str(column)
        normalized = normalize_token(column_text)
        column_tokens = {token for token in normalized.split("_") if len(token) > 1}
        sequence = SequenceMatcher(None, target, normalized).ratio() if target and normalized else 0.0
        union = target_tokens | column_tokens
        overlap = len(target_tokens & column_tokens) / len(union) if union else 0.0
        substring = 1.0 if target and (target in normalized or normalized in target) else 0.0
        preferred = 1.0 if preferred_column and normalize_token(preferred_column) == normalized else 0.0
        series = df[column]
        compatible = True
        if outcome_type == "continuous":
            compatible = pd.api.types.is_numeric_dtype(series) and not pd.api.types.is_bool_dtype(series)
        elif outcome_type in {"binary", "paired_binary", "diagnostic"}:
            compatible = _is_binary_series(series)
        elif outcome_type == "ordinal":
            compatible = _is_low_cardinality_series(series)
        elif outcome_type == "time_to_event":
            compatible = pd.api.types.is_numeric_dtype(series) or any(
                word in normalized for word in ("tiempo", "time", "dias", "days", "seguimiento", "followup", "evento")
            )
        type_bonus = 0.12 if compatible else -0.12
        clinical_bonus = 0.0
        clinical_aliases = {
            "presion_arterial_sistolica": ("pas", "sbp"),
            "presion_arterial_diastolica": ("pad", "dbp"),
            "frecuencia_cardiaca": ("fc", "hr"),
            "colesterol_ldl": ("ldl",),
            "hemoglobina_glicosilada": ("hba1c",),
        }
        for phrase, abbreviations in clinical_aliases.items():
            if phrase in target and any(abbreviation in normalized.split("_") or normalized.startswith(abbreviation) for abbreviation in abbreviations):
                clinical_bonus = 0.28
                break
        if preferred:
            score = 0.99 if compatible else 0.82
        elif target and target == normalized:
            score = 0.97 if compatible else 0.80
        else:
            score = min(max(0.44 * sequence + 0.30 * overlap + 0.08 * substring + clinical_bonus + type_bonus, 0.0), 1.0)
        candidates.append(
            {
                "column": column_text,
                "score": score,
                "compatible": compatible,
                "dtype": str(series.dtype),
                "unique": int(series.nunique(dropna=True)),
            }
        )
    return sorted(candidates, key=lambda item: (item["score"], item["compatible"]), reverse=True)


def confidence_label(score: float) -> str:
    if score >= 0.85:
        return "Alta"
    if score >= 0.65:
        return "Moderada"
    if score >= 0.45:
        return "Baja"
    return "Muy baja"


def invalidate_outcome_confirmation() -> None:
    st.session_state.outcome_confirmed = False


def extract_protocol_mapping_from_text(text: str) -> dict[str, Any]:
    cleaned = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"\s+", " ", line).strip() for line in cleaned.split("\n") if line.strip()]
    protocol: dict[str, Any] = {}

    for field_name, aliases in PROTOCOL_FIELD_ALIASES.items():
        raw_value = _value_after_alias(lines, aliases)
        if raw_value is None:
            continue
        if field_name in TEXT_PROTOCOL_FIELDS:
            protocol[field_name] = _clean_outcome_value(raw_value) if field_name == "outcome_primario" else str(raw_value).strip()
        else:
            try:
                protocol[field_name] = _parse_locale_number(raw_value)
            except ValueError:
                continue

    if "tipo_diseno" not in protocol:
        inferred = infer_design_from_text(cleaned)
        if inferred:
            protocol["tipo_diseno"] = inferred

    # Parámetros específicos pueden completar los cinco campos mínimos.
    design_code: str | None = None
    if protocol.get("tipo_diseno") is not None:
        try:
            design_code = resolve_design(protocol["tipo_diseno"])
        except ValueError:
            inferred = infer_design_from_text(str(protocol["tipo_diseno"]))
            if inferred:
                protocol["tipo_diseno"] = inferred
                design_code = inferred

    if "efecto_esperado" not in protocol:
        effect_candidates: dict[str, tuple[str, ...]] = {
            "diagnostic_accuracy": ("sensibilidad", "especificidad"),
            "odds_ratio": ("odds_ratio",),
            "risk_ratio": ("riesgo_relativo",),
            "survival": ("hazard_ratio",),
            "noninferiority_continuous": ("margen",),
            "noninferiority_binary": ("margen",),
            "equivalence_continuous": ("margen",),
            "equivalence_binary": ("margen",),
        }
        for key in effect_candidates.get(design_code or "", ("margen",)):
            if key in protocol:
                protocol["efecto_esperado"] = protocol[key]
                break

    if "variabilidad_estimada" not in protocol:
        variability_candidates: dict[str, tuple[str, ...]] = {
            "diagnostic_accuracy": ("prevalencia",),
            "odds_ratio": ("proporcion_control", "proporcion_grupo_2"),
            "risk_ratio": ("proporcion_control", "proporcion_grupo_2"),
            "survival": ("mortalidad_control",),
            "proportion_estimation": ("precision", "proporcion_esperada"),
            "noninferiority_binary": ("proporcion_grupo_2",),
            "equivalence_binary": ("proporcion_grupo_2",),
        }
        for key in variability_candidates.get(design_code or "", ()):
            if key in protocol:
                protocol["variabilidad_estimada"] = protocol[key]
                break

    return enrich_outcome_fields(protocol, cleaned)


def canonicalize_protocol_mapping(data: dict[str, Any]) -> dict[str, Any]:
    canonical: dict[str, Any] = {}
    alias_lookup: dict[str, str] = {}
    for canonical_name, aliases in PROTOCOL_FIELD_ALIASES.items():
        alias_lookup[normalize_token(canonical_name)] = canonical_name
        for alias in aliases:
            alias_lookup[normalize_token(alias)] = canonical_name

    for key, value in data.items():
        canonical_key = alias_lookup.get(normalize_token(key), normalize_token(key))
        canonical[canonical_key] = value
    return canonical


def parse_protocol(uploaded_file: Any) -> dict[str, Any]:
    raw = uploaded_file.getvalue()
    suffix = os.path.splitext(uploaded_file.name)[1].lower()
    extraction_warnings: list[str] = []
    extracted_text = ""

    if suffix == ".json":
        data = json.loads(raw.decode("utf-8-sig"))
    elif suffix in {".yaml", ".yml"}:
        data = yaml.safe_load(raw.decode("utf-8-sig"))
    elif suffix in {".pdf", ".docx", ".docm", ".doc"}:
        extracted_text, extraction_warnings = extract_protocol_text(raw, suffix)
        data = extract_protocol_mapping_from_text(extracted_text)
    else:
        raise ValueError("El protocolo debe ser JSON, YAML, PDF o Word (DOCX/DOCM/DOC).")

    if not isinstance(data, dict):
        raise ValueError("El archivo de protocolo debe contener parámetros identificables.")
    data = canonicalize_protocol_mapping(data)

    structured_outcomes = extract_structured_outcomes(data)
    primary_outcome = choose_primary_outcome(structured_outcomes)
    if primary_outcome:
        for key in (
            "outcome_primario",
            "definicion_outcome",
            "tipo_outcome",
            "unidad_outcome",
            "momento_evaluacion",
            "columna_dataset_outcome",
            "valor_evento",
        ):
            if data.get(key) in (None, "") and primary_outcome.get(key) not in (None, ""):
                data[key] = primary_outcome[key]
        data["_outcomes"] = structured_outcomes

    data = enrich_outcome_fields(data, extracted_text)

    # Convierte campos numéricos aun cuando JSON/YAML los hayan guardado como texto con coma o porcentaje.
    for key in list(data):
        if key in NUMERIC_PROTOCOL_FIELDS and data[key] not in (None, ""):
            data[key] = _parse_locale_number(data[key])

    if data.get("tipo_diseno") is None and extracted_text:
        inferred = infer_design_from_text(extracted_text)
        if inferred:
            data["tipo_diseno"] = inferred

    required = {"tipo_diseno", "error_alfa", "poder_estadistico", "efecto_esperado", "variabilidad_estimada"}
    missing = sorted(key for key in required if key not in data or data[key] in (None, ""))
    if missing:
        readable = ", ".join(missing)
        raise ValueError(
            "No fue posible identificar todos los parámetros obligatorios. Faltan: "
            + readable
            + ". En PDF/Word use etiquetas claras, por ejemplo: 'Error alfa: 0,05'."
        )

    data["_source_format"] = suffix.lstrip(".").upper()
    data["_extraction_warnings"] = extraction_warnings
    if extracted_text:
        data["_text_preview"] = extracted_text[:1500]
    return data

def resolve_design(value: Any) -> str:
    token = normalize_token(value)
    if token in DESIGNS:
        return token
    if token in ALIASES:
        return ALIASES[token]
    # Búsqueda tolerante por palabras del rótulo.
    for code, label in DESIGNS.items():
        normalized_label = normalize_token(label)
        if token == normalized_label or token in normalized_label or normalized_label in token:
            return code
    inferred = infer_design_from_text(str(value))
    if inferred:
        return inferred
    raise ValueError(f"tipo_diseno no reconocido: {value!r}")


def apply_protocol_to_state(protocol: dict[str, Any]) -> None:
    code = resolve_design(protocol["tipo_diseno"])
    alpha = float(protocol["error_alfa"])
    power = float(protocol["poder_estadistico"])
    if alpha > 1:
        alpha /= 100
    if power > 1:
        power /= 100
    if not 0 < alpha < 1:
        raise ValueError("error_alfa debe expresarse como proporción (0,05) o porcentaje (5).")
    if not 0.50 <= power < 1:
        raise ValueError("poder_estadistico debe estar entre 0,50 y menos de 1.")

    effect = float(protocol["efecto_esperado"])
    variability = float(protocol["variabilidad_estimada"])
    if not math.isfinite(effect) or not math.isfinite(variability):
        raise ValueError("efecto_esperado y variabilidad_estimada deben ser valores numéricos finitos.")

    def bounded_probability(value: float, fallback: float = 0.50) -> float:
        numeric = float(value)
        if 1 < numeric <= 100:
            numeric /= 100
        if 0 <= numeric <= 1:
            return min(max(numeric, 0.001), 0.999)
        return fallback

    def probability_or_points(value: float) -> float:
        numeric = abs(float(value))
        return numeric / 100 if 1 < numeric <= 100 else numeric

    st.session_state.design_code = code
    st.session_state.alpha = alpha
    st.session_state.power = power
    st.session_state.outcome_primary = str(protocol.get("outcome_primario", "") or "")
    st.session_state.outcome_definition = str(protocol.get("definicion_outcome", "") or "")
    st.session_state.outcome_type = infer_outcome_type(protocol.get("tipo_outcome", ""), code)
    st.session_state.outcome_unit = str(protocol.get("unidad_outcome", "") or "")
    st.session_state.outcome_timepoint = str(protocol.get("momento_evaluacion", "") or "")
    st.session_state.outcome_secondary = str(protocol.get("outcome_secundarios", "") or "")
    st.session_state.outcome_preferred_column = str(protocol.get("columna_dataset_outcome", "") or "")
    st.session_state.outcome_event_value_protocol = str(protocol.get("valor_evento", "") or "")
    st.session_state.protocol_outcomes = protocol.get("_outcomes", [])
    st.session_state.outcome_confirmed = False
    st.session_state.outcome_column_confidence = 0.0

    # Interpretación automática de los dos parámetros genéricos según el diseño.
    if code == "mean_estimation":
        st.session_state.precision_mean = max(abs(effect), 0.0001)
        st.session_state.sd = max(abs(variability), 0.0001)
    elif code in {"means_independent", "means_paired", "superiority_continuous"}:
        st.session_state.delta = max(abs(effect), 0.0001)
        st.session_state.sd = max(abs(variability), 0.0001)
        st.session_state.sd_diff = max(abs(variability), 0.0001)
    elif code == "proportion_estimation":
        st.session_state.p = bounded_probability(effect)
        st.session_state.precision_prop = max(min(probability_or_points(variability), 0.50), 0.001)
    elif code in {"proportions_independent", "superiority_binary"}:
        baseline = bounded_probability(variability)
        absolute_effect = probability_or_points(effect)
        st.session_state.p2 = baseline
        st.session_state.p1 = min(max(baseline + absolute_effect, 0.001), 0.999)
    elif code in {"noninferiority_continuous", "equivalence_continuous"}:
        st.session_state.margin_cont = max(abs(effect), 0.0001)
        st.session_state.sd = max(abs(variability), 0.0001)
    elif code == "noninferiority_binary":
        baseline = bounded_probability(variability, 0.80)
        st.session_state.margin_bin = max(min(probability_or_points(effect), 0.50), 0.001)
        st.session_state.p1_ni = baseline
        st.session_state.p2_ni = baseline
    elif code == "equivalence_binary":
        baseline = bounded_probability(variability, 0.80)
        st.session_state.margin_bin = max(min(probability_or_points(effect), 0.50), 0.001)
        st.session_state.p1_eq = baseline
        st.session_state.p2_eq = baseline
    elif code == "mcnemar":
        discordance_total = min(max(probability_or_points(variability), 0.01), 0.98)
        discordance_difference = min(probability_or_points(effect), discordance_total - 0.001)
        st.session_state.p01 = (discordance_total + discordance_difference) / 2
        st.session_state.p10 = (discordance_total - discordance_difference) / 2
    elif code in {"odds_ratio", "risk_ratio"}:
        st.session_state.effect_ratio = max(abs(effect), 0.01)
        st.session_state.p_control = bounded_probability(variability, 0.20)
    elif code == "diagnostic_accuracy":
        accuracy = bounded_probability(effect, 0.90)
        st.session_state.sensitivity = accuracy
        st.session_state.specificity = accuracy
        st.session_state.prevalence = bounded_probability(variability, 0.20)
    elif code == "survival":
        hr = max(abs(effect), 0.01)
        mortality_control = bounded_probability(variability, 0.30)
        st.session_state.hazard_ratio = hr
        st.session_state.mortality_control = mortality_control
        st.session_state.mortality_treatment = min(max(mortality_control * hr, 0.001), 0.999)

    if "precision" in protocol and protocol["precision"] is not None:
        precision_value = abs(float(protocol["precision"]))
        if code == "mean_estimation":
            st.session_state.precision_mean = max(precision_value, 0.0001)
        elif code == "proportion_estimation":
            st.session_state.precision_prop = max(min(probability_or_points(precision_value), 0.50), 0.001)
    if "margen" in protocol and protocol["margen"] is not None:
        margin_value = abs(float(protocol["margen"]))
        if code in {"noninferiority_continuous", "equivalence_continuous"}:
            st.session_state.margin_cont = max(margin_value, 0.0001)
        elif code in {"noninferiority_binary", "equivalence_binary"}:
            st.session_state.margin_bin = max(min(probability_or_points(margin_value), 0.50), 0.001)
    if "perdidas_esperadas" in protocol and protocol["perdidas_esperadas"] is not None:
        loss_value = float(protocol["perdidas_esperadas"])
        loss_percent = loss_value * 100 if loss_value <= 1 else loss_value
        st.session_state.loss_percent = int(round(min(max(loss_percent, 0), 50)))

    # Los campos específicos, cuando existen, prevalecen sobre la interpretación genérica.
    if "proporcion_grupo_1" in protocol and protocol["proporcion_grupo_1"] is not None:
        target = "p1_eq" if code == "equivalence_binary" else "p1_ni" if code == "noninferiority_binary" else "p1"
        st.session_state[target] = bounded_probability(protocol["proporcion_grupo_1"])
    if "proporcion_grupo_2" in protocol and protocol["proporcion_grupo_2"] is not None:
        target = "p2_eq" if code == "equivalence_binary" else "p2_ni" if code == "noninferiority_binary" else "p2"
        st.session_state[target] = bounded_probability(protocol["proporcion_grupo_2"])

    optional_map = {
        "proporcion_esperada": "p",
        "proporcion_control": "p_control",
        "diferencia_esperada": "expected_diff",
        "prevalencia": "prevalence",
        "sensibilidad": "sensitivity",
        "especificidad": "specificity",
        "precision_sensibilidad": "precision_se",
        "precision_especificidad": "precision_sp",
        "odds_ratio": "effect_ratio",
        "riesgo_relativo": "effect_ratio",
        "hazard_ratio": "hazard_ratio",
        "mortalidad_control": "mortality_control",
        "mortalidad_tratamiento": "mortality_treatment",
        "poblacion": "population",
        "p01": "p01",
        "p10": "p10",
    }
    probability_fields = {
        "proporcion_esperada",
        "proporcion_control",
        "prevalencia",
        "sensibilidad",
        "especificidad",
        "precision_sensibilidad",
        "precision_especificidad",
        "mortalidad_control",
        "mortalidad_tratamiento",
        "p01",
        "p10",
    }
    for source_key, state_key in optional_map.items():
        if source_key in protocol and protocol[source_key] is not None:
            value = protocol[source_key]
            if source_key in probability_fields:
                value = bounded_probability(value)
            st.session_state[state_key] = value


def load_dataset(uploaded_file: Any) -> tuple[pd.DataFrame | None, dict[str, Any]]:
    suffix = os.path.splitext(uploaded_file.name)[1].lower()
    metadata: dict[str, Any] = {"name": uploaded_file.name, "type": suffix}
    if suffix == ".csv":
        raw = uploaded_file.getvalue()
        last_error: Exception | None = None
        for encoding in ("utf-8-sig", "latin-1"):
            for sep in (None, ",", ";", "\t"):
                try:
                    df = pd.read_csv(io.BytesIO(raw), encoding=encoding, sep=sep, engine="python")
                    if df.shape[1] > 0:
                        metadata["rows"] = len(df)
                        metadata["columns"] = len(df.columns)
                        return df, metadata
                except Exception as exc:  # pragma: no cover - depende del archivo
                    last_error = exc
        raise ValueError(f"No se pudo interpretar el CSV: {last_error}")

    if suffix in {".xlsx", ".xlsm"}:
        xls = pd.ExcelFile(io.BytesIO(uploaded_file.getvalue()), engine="openpyxl")
        sheet = st.sidebar.selectbox("Hoja de Excel", xls.sheet_names, key="excel_sheet")
        df = pd.read_excel(xls, sheet_name=sheet, engine="openpyxl")
        metadata.update({"sheet": sheet, "rows": len(df), "columns": len(df.columns)})
        return df, metadata

    if suffix in {".sqlite", ".sqlite3", ".db"}:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        tmp.write(uploaded_file.getvalue())
        tmp.close()
        engine = create_engine(f"sqlite+pysqlite:///{tmp.name}")
        table_names = inspect(engine).get_table_names()
        if not table_names:
            raise ValueError("La base SQLite no contiene tablas visibles.")
        table = st.sidebar.selectbox("Tabla SQLite", table_names, key="sqlite_table")
        df = pd.read_sql_table(table, con=engine)
        metadata.update({"table": table, "rows": len(df), "columns": len(df.columns), "temp_path": tmp.name})
        return df, metadata

    raise ValueError("Formato no compatible. Use XLSX, CSV, SQLite, SQLite3 o DB.")


def numeric_columns(df: pd.DataFrame) -> list[str]:
    return [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]


# -----------------------------------------------------------------------------
# Reportes DOCX y anexos CONSORT
# -----------------------------------------------------------------------------
CONSORT_ITEMS = [
    ("Título y resumen", "Identificar el estudio como aleatorizado y presentar un resumen estructurado."),
    ("Fundamento", "Explicar el contexto científico y los objetivos o hipótesis."),
    ("Diseño", "Describir el diseño, la razón de asignación y cualquier cambio posterior al inicio."),
    ("Participantes", "Precisar criterios de elegibilidad, centros y lugares de recolección."),
    ("Intervenciones", "Detallar cada intervención con información suficiente para reproducirla."),
    ("Desenlaces", "Definir desenlaces primarios/secundarios y cómo/cuándo se evaluaron."),
    ("Tamaño de muestra", "Informar cómo se determinó la muestra y cualquier análisis intermedio o regla de detención."),
    ("Secuencia aleatoria", "Describir el método de generación de la secuencia y restricciones."),
    ("Ocultamiento", "Explicar el mecanismo utilizado para ocultar la asignación."),
    ("Implementación", "Indicar quién generó la secuencia, incorporó participantes y asignó intervenciones."),
    ("Enmascaramiento", "Informar quién estuvo cegado y cómo se mantuvo el cegamiento."),
    ("Métodos estadísticos", "Describir análisis primarios, secundarios y métodos adicionales."),
    ("Flujo de participantes", "Documentar asignación, seguimiento, pérdidas, exclusiones y análisis."),
    ("Reclutamiento", "Informar fechas de reclutamiento y seguimiento, y motivo de finalización."),
    ("Datos basales", "Presentar características demográficas y clínicas por grupo."),
    ("Números analizados", "Indicar denominadores y si se analizó según asignación original."),
    ("Resultados y estimación", "Informar efecto, precisión e intervalos de confianza."),
    ("Análisis adicionales", "Distinguir análisis preespecificados de exploratorios."),
    ("Daños", "Describir eventos adversos o efectos no deseados por grupo."),
    ("Limitaciones", "Discutir sesgos potenciales, imprecisión y multiplicidad."),
    ("Generalización", "Analizar la aplicabilidad externa de los hallazgos."),
    ("Interpretación", "Ofrecer una interpretación consistente con beneficios, daños y evidencia disponible."),
    ("Registro", "Indicar registro del ensayo y número de inscripción."),
    ("Protocolo", "Señalar dónde puede accederse al protocolo completo."),
    ("Financiación", "Informar fuentes de financiación, apoyos y papel de los financiadores."),
]


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(12)


def add_key_value_table(doc: Document, rows: list[tuple[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Parámetro"
    table.rows[0].cells[1].text = "Valor"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)


def consort_paragraph(
    result: SampleSizeResult,
    params: dict[str, Any],
    recommended_total: int | None = None,
) -> str:
    alpha = float(params.get("alpha", 0.05))
    power = float(params.get("power", 0.80))
    sided = params.get("sided", "Bilateral").lower()
    losses = float(params.get("loss_rate", 0.0))
    variability = params.get("sd", params.get("sd_diff", "no aplicable"))
    effect = params.get("delta", params.get("margin", params.get("effect_ratio", "según diseño")))
    outcome = params.get("outcome", {}) or {}
    outcome_name = outcome.get("outcome_primario") or "desenlace primario preespecificado"
    outcome_type = outcome.get("tipo_outcome_label") or outcome.get("tipo_outcome") or "no especificado"
    unit = outcome.get("unidad_outcome") or "sin unidad aplicable"
    timepoint = outcome.get("momento_evaluacion") or "momento preespecificado"
    power_clause = (
        f"un poder estadístico de {power:.0%}" if result.uses_power else "un criterio de precisión del intervalo de confianza"
    )
    total_clause = result.n_final_total
    multiplicity_clause = ""
    if recommended_total and recommended_total > result.n_final_total:
        total_clause = recommended_total
        multiplicity_clause = (
            f" Debido a la existencia de outcomes coprimarios con parámetros completos, se adoptó el mayor requerimiento, "
            f"correspondiente a {recommended_total} participantes."
        )
    return (
        f"El tamaño de muestra se determinó a priori para el outcome primario '{outcome_name}', de tipo {str(outcome_type).lower()}, "
        f"medido en {unit} y evaluado en {timepoint}, mediante {result.method.lower()} "
        f"Se asumió un error alfa de {alpha:.3f}, un contraste {sided}, {power_clause}, "
        f"un efecto o margen de {effect} y una variabilidad/proporción basal de {variability}. "
        f"El cálculo inicial fue de {math.ceil(result.n_raw_total)} participantes y se ajustó por una pérdida esperada "
        f"de {losses:.1%}, obteniéndose un requerimiento final de {total_clause}."
        f"{multiplicity_clause} Estos parámetros, la definición operacional del outcome, su momento de evaluación, "
        "la fuente del efecto clínicamente relevante y cualquier modificación posterior deberán declararse en Métodos, "
        "de acuerdo con el ítem de tamaño de muestra de CONSORT 2010."
    )


def build_word_report(
    result: SampleSizeResult,
    params: dict[str, Any],
    sensitivity_df: pd.DataFrame,
    dataset_metadata: dict[str, Any] | None,
    coprimary_df: pd.DataFrame | None = None,
    recommended_total: int | None = None,
) -> bytes:
    doc = Document()
    set_doc_defaults(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Informe de cálculo del tamaño de muestra")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{APP_NAME} · v{APP_VERSION} · {datetime.now():%d/%m/%Y %H:%M}")

    outcome = params.get("outcome", {}) or {}
    final_recommendation = int(recommended_total or result.n_final_total)

    doc.add_heading("1. Resumen del diseño", level=1)
    doc.add_paragraph(result.design_label)
    add_key_value_table(
        doc,
        [
            ("N total recomendado", final_recommendation),
            ("N del outcome primario confirmado", result.n_final_total),
            ("N inicial sin pérdidas", math.ceil(result.n_raw_total)),
            ("Método", result.method),
            ("Distribución", "; ".join(f"{k}: {v}" for k, v in result.groups_final.items())),
        ],
    )

    doc.add_heading("2. Outcome primario y trazabilidad", level=1)
    add_key_value_table(
        doc,
        [
            ("Outcome primario", outcome.get("outcome_primario", "")),
            ("Definición operacional", outcome.get("definicion_outcome", "")),
            ("Tipo", outcome.get("tipo_outcome_label", outcome.get("tipo_outcome", ""))),
            ("Unidad", outcome.get("unidad_outcome", "")),
            ("Momento de evaluación", outcome.get("momento_evaluacion", "")),
            ("Columna del dataset", outcome.get("columna_dataset_outcome", "No aplicada")),
            ("Confianza de vinculación", f"{float(outcome.get('confianza_columna', 0.0)):.0%}" if dataset_metadata else "No aplicada"),
            ("Valor del evento", outcome.get("valor_evento", "")),
            ("Confirmado por el investigador", "Sí" if outcome.get("confirmado") else "No"),
        ],
    )
    if outcome.get("outcome_secundarios"):
        p = doc.add_paragraph()
        p.add_run("Outcomes secundarios/coprimarios declarados: ").bold = True
        p.add_run(str(outcome.get("outcome_secundarios")))

    doc.add_heading("3. Justificación matemática", level=1)
    doc.add_paragraph(result.formula)
    doc.add_paragraph("Supuestos utilizados: " + "; ".join(result.assumptions) + ".")
    if result.warnings:
        p = doc.add_paragraph()
        p.add_run("Advertencias metodológicas: ").bold = True
        p.add_run(" ".join(result.warnings))

    if coprimary_df is not None and not coprimary_df.empty:
        doc.add_heading("4. Evaluación de outcomes primarios/coprimarios", level=1)
        doc.add_paragraph(
            "Cuando el protocolo aportó parámetros completos para más de un outcome primario o coprimario, se calculó cada escenario por separado. "
            "La recomendación global adopta el mayor tamaño de muestra para evitar que uno de los outcomes quede insuficientemente dimensionado."
        )
        table = doc.add_table(rows=1, cols=len(coprimary_df.columns))
        table.style = "Table Grid"
        for idx, column in enumerate(coprimary_df.columns):
            table.rows[0].cells[idx].text = str(column)
        for _, row in coprimary_df.iterrows():
            cells = table.add_row().cells
            for idx, value in enumerate(row):
                cells[idx].text = str(value)
        next_section = 5
    else:
        next_section = 4

    doc.add_heading(f"{next_section}. Diferencia mínima clínicamente importante", level=1)
    doc.add_paragraph(
        "El efecto esperado debe representar la diferencia mínima que modificaría una decisión clínica, no la diferencia "
        "que simplemente podría alcanzar significación estadística. Su valor debe justificarse con literatura previa, "
        "datos piloto o consenso clínico-estadístico."
    )

    doc.add_heading(f"{next_section + 1}. Párrafo para reporte CONSORT 2010", level=1)
    doc.add_paragraph(consort_paragraph(result, params, final_recommendation))

    doc.add_heading(f"{next_section + 2}. Análisis de sensibilidad", level=1)
    table = doc.add_table(rows=1, cols=len(sensitivity_df.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(sensitivity_df.columns):
        table.rows[0].cells[idx].text = str(column)
    for _, row in sensitivity_df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)

    section_number = next_section + 3
    if dataset_metadata:
        doc.add_heading(f"{section_number}. Fuente de datos importada", level=1)
        doc.add_paragraph(
            "; ".join(f"{k}: {v}" for k, v in dataset_metadata.items() if k != "temp_path")
        )
        section_number += 1

    doc.add_heading(f"{section_number}. Trazabilidad y límites", level=1)
    doc.add_paragraph(
        "Este informe documenta un cálculo de planificación. La detección automática del outcome y de su columna es una propuesta asistida y no reemplaza la confirmación clínica. "
        "Los diseños con múltiples desenlaces, análisis intermedios, conglomerados, estratificación, multiplicidad, riesgos no proporcionales o requisitos regulatorios "
        "deben validarse con un bioestadístico y, cuando corresponda, mediante simulación."
    )

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_consort_docx() -> bytes:
    doc = Document()
    set_doc_defaults(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Checklist operativo CONSORT 2010")
    doc.add_paragraph(
        "Versión para trabajo interno, redactada en forma resumida y parafraseada. Debe cotejarse con el documento oficial."
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["N.º", "Dominio", "Verificación", "Página/nota"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for i, (domain, check) in enumerate(CONSORT_ITEMS, 1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = domain
        cells[2].text = check
        cells[3].text = ""
    doc.add_paragraph(
        "Fuente normativa: CONSORT 2010 Statement y checklist oficial. Para ensayos nuevos, verificar además las actualizaciones vigentes."
    )
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_consort_pdf() -> bytes:
    bio = io.BytesIO()
    doc = SimpleDocTemplate(
        bio,
        pagesize=A4,
        rightMargin=1.2 * cm,
        leftMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitle", parent=styles["Title"], alignment=TA_CENTER, fontSize=16, leading=19))
    story = [Paragraph("Checklist operativo CONSORT 2010", styles["CenterTitle"]), Spacer(1, 0.25 * cm)]
    story.append(
        Paragraph(
            "Versión resumida y parafraseada para trabajo interno; cotejar con el checklist oficial.",
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 0.25 * cm))
    data = [["N.º", "Dominio", "Verificación", "Página/nota"]]
    for i, (domain, check) in enumerate(CONSORT_ITEMS, 1):
        data.append([str(i), domain, Paragraph(check, styles["BodyText"]), ""])
    table = Table(data, colWidths=[1.0 * cm, 3.5 * cm, 11.0 * cm, 2.2 * cm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCEAF7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#12344D")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.2),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7FAFC")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.25 * cm))
    story.append(
        Paragraph(
            "Fuente normativa: CONSORT 2010 Statement. Para ensayos nuevos, verificar además las actualizaciones vigentes.",
            styles["BodyText"],
        )
    )
    doc.build(story)
    return bio.getvalue()



def _protocol_display_rows(protocol: dict[str, Any]) -> list[tuple[str, str]]:
    labels = {
        "tipo_diseno": "Tipo de diseño",
        "error_alfa": "Error alfa",
        "poder_estadistico": "Poder estadístico",
        "efecto_esperado": "Efecto esperado",
        "variabilidad_estimada": "Variabilidad estimada",
        "perdidas_esperadas": "Pérdidas esperadas",
        "outcome_primario": "Desenlace primario",
        "definicion_outcome": "Definición del desenlace",
        "tipo_outcome": "Tipo de outcome",
        "unidad_outcome": "Unidad del outcome",
        "momento_evaluacion": "Momento de evaluación",
        "columna_dataset_outcome": "Columna del dataset",
        "valor_evento": "Valor del evento",
        "outcome_secundarios": "Desenlaces secundarios",
    }
    return [(labels[key], str(protocol[key])) for key in labels if key in protocol]


def build_protocol_docx_template(protocol: dict[str, Any]) -> bytes:
    document = Document()
    document.add_heading("Archivo de configuración de protocolo", level=1)
    document.add_paragraph(
        "Complete los valores conservando las etiquetas. BioSize Clinical leerá automáticamente este archivo Word."
    )
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Parámetro"
    table.rows[0].cells[1].text = "Valor"
    for label, value in _protocol_display_rows(protocol):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.add_paragraph(
        "También pueden agregarse parámetros específicos, por ejemplo: prevalencia, sensibilidad, especificidad, margen, hazard ratio o proporciones por grupo."
    )
    bio = io.BytesIO()
    document.save(bio)
    return bio.getvalue()


def build_protocol_pdf_template(protocol: dict[str, Any]) -> bytes:
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    story: list[Any] = [
        Paragraph("Archivo de configuración de protocolo", styles["Title"]),
        Spacer(1, 0.3 * cm),
        Paragraph(
            "Complete los valores conservando las etiquetas. BioSize Clinical leerá automáticamente este PDF cuando contenga texto seleccionable.",
            styles["BodyText"],
        ),
        Spacer(1, 0.4 * cm),
    ]
    rows = [["Parámetro", "Valor"]] + [[label, value] for label, value in _protocol_display_rows(protocol)]
    table = Table(rows, colWidths=[9 * cm, 5 * cm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbe7f0")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)
    doc.build(story)
    return bio.getvalue()


# -----------------------------------------------------------------------------
# Interfaz Streamlit
# -----------------------------------------------------------------------------
st.set_page_config(page_title=APP_NAME, page_icon="📐", layout="wide")

st.markdown(
    """
    <style>
      .main .block-container {padding-top: 1.5rem; padding-bottom: 3rem; max-width: 1450px;}
      .hero {padding: 1.2rem 1.4rem; border: 1px solid #dbe7f0; border-radius: 18px;
             background: linear-gradient(135deg, #f6fbff 0%, #eef7f4 100%); margin-bottom: 1rem;}
      .hero h1 {margin: 0; color: #12344d; font-size: 2.2rem;}
      .hero p {margin: .35rem 0 0 0; color: #496273;}
      .metric-card {border: 1px solid #dbe7f0; border-radius: 16px; padding: 1rem;
                    background: #ffffff; min-height: 120px; box-shadow: 0 2px 10px rgba(22,52,74,.05);}
      .metric-label {font-size: .86rem; color: #627787; text-transform: uppercase; letter-spacing: .04em;}
      .metric-value {font-size: 2.15rem; font-weight: 750; color: #0f5e62; margin-top: .15rem;}
      .metric-note {font-size: .86rem; color: #6f7f8b;}
      .technical-note {border-left: 4px solid #0f7b78; padding: .65rem .9rem; background: #f1fbfa; border-radius: 8px;}
      div[data-testid="stDownloadButton"] button {border-radius: 10px;}
    </style>
    """,
    unsafe_allow_html=True,
)

# Estado inicial
DEFAULTS = {
    "design_code": "means_independent",
    "alpha": 0.05,
    "power": 0.80,
    "sided": "Bilateral",
    "loss_percent": 10,
    "delta": 5.0,
    "sd": 10.0,
    "sd_diff": 8.0,
    "precision_mean": 2.0,
    "precision_prop": 0.05,
    "p": 0.50,
    "p1": 0.60,
    "p2": 0.40,
    "p1_ni": 0.78,
    "p2_ni": 0.80,
    "p1_eq": 0.80,
    "p2_eq": 0.80,
    "p01": 0.20,
    "p10": 0.08,
    "margin_cont": 5.0,
    "margin_bin": 0.10,
    "expected_diff": 0.0,
    "p_control": 0.20,
    "effect_ratio": 2.0,
    "sensitivity": 0.90,
    "specificity": 0.90,
    "prevalence": 0.20,
    "precision_se": 0.05,
    "precision_sp": 0.05,
    "hazard_ratio": 0.70,
    "mortality_control": 0.30,
    "mortality_treatment": 0.22,
    "allocation_treatment": 0.50,
    "finite_population": False,
    "population": 1000,
    "yates": False,
    "outcome_primary": "",
    "outcome_definition": "",
    "outcome_type": "unknown",
    "outcome_unit": "",
    "outcome_timepoint": "",
    "outcome_secondary": "",
    "outcome_preferred_column": "",
    "outcome_dataset_column": "",
    "outcome_event_value_protocol": "",
    "outcome_confirmed": False,
    "outcome_column_confidence": 0.0,
    "protocol_outcomes": [],
}
for key, value in DEFAULTS.items():
    st.session_state.setdefault(key, value)

with st.sidebar:
    st.header("📥 Importación")
    protocol_file = st.file_uploader(
        "Archivo de protocolo",
        type=["json", "yaml", "yml", "pdf", "docx", "docm", "doc"],
        key="protocol_upload",
        help="PDF y Word deben contener texto seleccionable o una tabla con los parámetros del protocolo.",
    )
    if protocol_file is not None:
        signature = (protocol_file.name, len(protocol_file.getvalue()), hash(protocol_file.getvalue()))
        if st.session_state.get("protocol_signature") != signature:
            try:
                protocol = parse_protocol(protocol_file)
                apply_protocol_to_state(protocol)
                st.session_state.protocol_signature = signature
                st.session_state.protocol_loaded_message = (
                    f"Protocolo aplicado: {protocol_file.name} · formato {protocol.get('_source_format', '')}"
                )
                st.session_state.protocol_detected = {
                    key: value for key, value in protocol.items() if not key.startswith("_")
                }
                st.session_state.protocol_extraction_warnings = protocol.get("_extraction_warnings", [])
                st.session_state.protocol_text_preview = protocol.get("_text_preview", "")
            except Exception as exc:
                st.error(str(exc))
    if st.session_state.get("protocol_loaded_message"):
        st.success(st.session_state.protocol_loaded_message)
    for warning in st.session_state.get("protocol_extraction_warnings", []):
        st.warning(warning)
    if st.session_state.get("protocol_detected"):
        with st.expander("Parámetros detectados", expanded=False):
            st.json(st.session_state.protocol_detected)
    if st.session_state.get("protocol_text_preview"):
        with st.expander("Vista previa del texto extraído", expanded=False):
            st.text(st.session_state.protocol_text_preview)

    data_file = st.file_uploader("Dataset", type=["xlsx", "xlsm", "csv", "sqlite", "sqlite3", "db"], key="data_upload")
    dataset: pd.DataFrame | None = None
    dataset_metadata: dict[str, Any] | None = None
    if data_file is not None:
        try:
            dataset, dataset_metadata = load_dataset(data_file)
            dataset_signature = (data_file.name, len(data_file.getvalue()), hash(data_file.getvalue()))
            if st.session_state.get("dataset_signature") != dataset_signature:
                st.session_state.dataset_signature = dataset_signature
                st.session_state.outcome_confirmed = False
                st.session_state.outcome_column_confidence = 0.0
            st.success(f"{len(dataset):,} filas · {len(dataset.columns)} columnas")
        except Exception as exc:
            st.error(f"No se pudo cargar el dataset: {exc}")

    st.divider()
    st.caption("Plantillas de protocolo")
    sample_protocol = {
        "tipo_diseno": "medias_independientes",
        "error_alfa": 0.05,
        "poder_estadistico": 0.80,
        "efecto_esperado": 5.0,
        "variabilidad_estimada": 10.0,
        "perdidas_esperadas": 0.10,
        "outcome_primario": "Presión arterial sistólica a las 12 semanas",
        "definicion_outcome": "Cambio de la presión arterial sistólica desde el valor basal hasta la semana 12",
        "tipo_outcome": "continuo",
        "unidad_outcome": "mmHg",
        "momento_evaluacion": "12 semanas",
        "columna_dataset_outcome": "PAS_12semanas",
        "outcome_secundarios": "Presión arterial diastólica y proporción de pacientes controlados",
    }
    sample_coprimary_protocol = {
        "tipo_diseno": "medias_independientes",
        "error_alfa": 0.05,
        "poder_estadistico": 0.80,
        "efecto_esperado": 5.0,
        "variabilidad_estimada": 10.0,
        "perdidas_esperadas": 0.10,
        "outcomes": [
            {
                "nombre": "Presión arterial sistólica a las 12 semanas",
                "rol": "primario",
                "tipo": "continuo",
                "tipo_diseno": "medias_independientes",
                "efecto_esperado": 5.0,
                "variabilidad_estimada": 10.0,
            },
            {
                "nombre": "Presión arterial diastólica a las 12 semanas",
                "rol": "coprimario",
                "tipo": "continuo",
                "tipo_diseno": "medias_independientes",
                "efecto_esperado": 3.0,
                "variabilidad_estimada": 10.0,
            },
        ],
    }
    st.download_button(
        "Descargar JSON de ejemplo",
        data=json.dumps(sample_protocol, indent=2, ensure_ascii=False).encode("utf-8"),
        file_name="protocolo_ejemplo.json",
        mime="application/json",
        use_container_width=True,
    )
    st.download_button(
        "Descargar YAML de ejemplo",
        data=yaml.safe_dump(sample_protocol, allow_unicode=True, sort_keys=False).encode("utf-8"),
        file_name="protocolo_ejemplo.yaml",
        mime="application/x-yaml",
        use_container_width=True,
    )
    st.download_button(
        "Descargar JSON con outcomes coprimarios",
        data=json.dumps(sample_coprimary_protocol, indent=2, ensure_ascii=False).encode("utf-8"),
        file_name="protocolo_coprimarios_ejemplo.json",
        mime="application/json",
        use_container_width=True,
    )
    try:
        st.download_button(
            "Descargar Word de ejemplo",
            data=build_protocol_docx_template(sample_protocol),
            file_name="protocolo_ejemplo.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.download_button(
            "Descargar PDF de ejemplo",
            data=build_protocol_pdf_template(sample_protocol),
            file_name="protocolo_ejemplo.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    except Exception as exc:
        st.caption(f"No se pudieron generar las plantillas documentales: {exc}")

st.markdown(
    """
    <div class="hero">
      <h1>📐 BioSize Clinical</h1>
      <p>Cálculo trazable de tamaño de muestra para investigación clínica, con curva de potencia, ajustes metodológicos y reporte Word editable.</p>
    </div>
    """,
    unsafe_allow_html=True,
)

# Selector de diseño y parámetros generales
left, right = st.columns([1.4, 1.0], gap="large")
with left:
    design_code = st.selectbox(
        "Diseño del estudio",
        options=list(DESIGNS.keys()),
        format_func=lambda x: DESIGNS[x],
        key="design_code",
        on_change=invalidate_outcome_confirmation,
    )
with right:
    no_variability = st.checkbox("No dispongo de una estimación confiable de variabilidad", key="no_variability")

if no_variability and design_code in {
    "mean_estimation",
    "means_independent",
    "means_paired",
    "superiority_continuous",
    "noninferiority_continuous",
    "equivalence_continuous",
}:
    st.warning(
        "Se recomienda realizar primero un estudio piloto de aproximadamente 30–50 sujetos para estimar la desviación estándar o la variabilidad de las diferencias."
    )


st.subheader("🎯 Desenlace primario")
st.caption(
    "La app propone el outcome a partir del protocolo, pero el investigador debe confirmar su definición y, cuando exista un dataset, la columna correspondiente antes de calcular."
)

outcome_top_1, outcome_top_2, outcome_top_3, outcome_top_4 = st.columns([1.7, 1.0, 0.8, 0.9])
with outcome_top_1:
    outcome_primary = st.text_input(
        "Nombre del outcome primario",
        key="outcome_primary",
        placeholder="Ej.: Presión arterial sistólica a las 12 semanas",
        on_change=invalidate_outcome_confirmation,
    )
with outcome_top_2:
    outcome_type = st.selectbox(
        "Tipo de outcome",
        options=list(OUTCOME_TYPE_LABELS.keys()),
        format_func=lambda code: OUTCOME_TYPE_LABELS[code],
        key="outcome_type",
        on_change=invalidate_outcome_confirmation,
    )
with outcome_top_3:
    outcome_unit = st.text_input(
        "Unidad",
        key="outcome_unit",
        placeholder="mmHg",
        on_change=invalidate_outcome_confirmation,
    )
with outcome_top_4:
    outcome_timepoint = st.text_input(
        "Momento de evaluación",
        key="outcome_timepoint",
        placeholder="12 semanas",
        on_change=invalidate_outcome_confirmation,
    )

outcome_definition = st.text_area(
    "Definición operacional del outcome",
    key="outcome_definition",
    placeholder="Defina exactamente qué se mide, cómo se calcula y en qué momento.",
    height=82,
    on_change=invalidate_outcome_confirmation,
)
outcome_secondary = st.text_area(
    "Outcomes secundarios o coprimarios declarados",
    key="outcome_secondary",
    placeholder="Opcional. Identifique explícitamente cuáles son coprimarios.",
    height=68,
    on_change=invalidate_outcome_confirmation,
)

if outcome_primary and outcome_type == "unknown":
    inferred_type = infer_outcome_type(outcome_primary + " " + outcome_definition, design_code)
    if inferred_type != "unknown":
        st.info(f"Tipo sugerido por el texto y el diseño: **{OUTCOME_TYPE_LABELS[inferred_type]}**.")
        if st.button("Aplicar tipo sugerido", key="apply_inferred_outcome_type"):
            st.session_state.outcome_type = inferred_type
            st.session_state.outcome_confirmed = False
            st.rerun()

if not design_compatible_with_outcome(design_code, outcome_type):
    suggested_design = suggested_design_for_outcome(outcome_type, design_code)
    st.warning(
        f"El diseño seleccionado ({DESIGNS[design_code]}) no es coherente con un outcome {OUTCOME_TYPE_LABELS.get(outcome_type, outcome_type).lower()}."
    )
    if suggested_design and st.button(
        f"Alinear con: {DESIGNS[suggested_design]}",
        key="align_design_with_outcome",
    ):
        st.session_state.design_code = suggested_design
        st.session_state.outcome_confirmed = False
        st.rerun()

selected_outcome_column = ""
outcome_column_score = 0.0
outcome_event_value: Any = st.session_state.get("outcome_event_value_protocol", "")
outcome_column_dtype = ""
outcome_column_unique = None
outcome_column_valid = True

if dataset is not None:
    st.markdown("#### Vinculación con el dataset")
    column_map = {str(column): column for column in dataset.columns}
    candidates = outcome_column_candidates(
        dataset,
        outcome_primary,
        outcome_type,
        st.session_state.get("outcome_preferred_column", ""),
    )
    ranked_names = [item["column"] for item in candidates]
    ordered_columns = list(dict.fromkeys(ranked_names + list(column_map.keys())))
    current_column = str(st.session_state.get("outcome_dataset_column", "") or "")
    preferred_column = str(st.session_state.get("outcome_preferred_column", "") or "")
    preferred_match = next(
        (name for name in ordered_columns if normalize_token(name) == normalize_token(preferred_column)),
        "",
    )
    if current_column not in ordered_columns:
        if preferred_match:
            st.session_state.outcome_dataset_column = preferred_match
        elif candidates and candidates[0]["score"] >= 0.35:
            st.session_state.outcome_dataset_column = candidates[0]["column"]
        elif ordered_columns:
            st.session_state.outcome_dataset_column = ordered_columns[0]

    link_c1, link_c2 = st.columns([1.5, 1.0])
    with link_c1:
        selected_outcome_column = st.selectbox(
            "Columna que representa el outcome primario",
            options=ordered_columns,
            key="outcome_dataset_column",
            on_change=invalidate_outcome_confirmation,
        )
    selected_candidate = next((item for item in candidates if item["column"] == selected_outcome_column), None)
    if selected_candidate:
        outcome_column_score = float(selected_candidate["score"])
        outcome_column_dtype = str(selected_candidate["dtype"])
        outcome_column_unique = int(selected_candidate["unique"])
        st.session_state.outcome_column_confidence = outcome_column_score
    with link_c2:
        st.metric("Confianza de correspondencia", f"{outcome_column_score:.0%}")
        st.caption(
            f"{confidence_label(outcome_column_score)} · tipo {outcome_column_dtype or 'no determinado'} · {outcome_column_unique if outcome_column_unique is not None else '—'} valores únicos"
        )

    if candidates:
        with st.expander("Alternativas de columna sugeridas", expanded=False):
            candidate_table = pd.DataFrame(
                [
                    {
                        "Columna": item["column"],
                        "Confianza": f"{item['score']:.0%}",
                        "Compatible con el tipo": "Sí" if item["compatible"] else "No",
                        "Tipo": item["dtype"],
                        "Valores únicos": item["unique"],
                    }
                    for item in candidates[:8]
                ]
            )
            st.dataframe(candidate_table, hide_index=True, use_container_width=True)

    if selected_outcome_column:
        actual_column = column_map[selected_outcome_column]
        outcome_series = dataset[actual_column].dropna()
        if outcome_type in {"binary", "paired_binary", "diagnostic"} and not outcome_series.empty:
            unique_values = list(outcome_series.unique())[:50]
            if len(unique_values) > 2:
                outcome_column_valid = False
                st.warning(
                    f"La columna seleccionada contiene {len(unique_values)} categorías. Un outcome binario requiere dos categorías válidas; recodifique o filtre la variable antes de usarla."
                )
            current_event = st.session_state.get("outcome_event_value")
            protocol_event = str(st.session_state.get("outcome_event_value_protocol", "") or "")
            if current_event not in unique_values:
                matched_event = next((value for value in unique_values if str(value) == protocol_event), None)
                st.session_state.outcome_event_value = matched_event if matched_event is not None else unique_values[-1]
            outcome_event_value = st.selectbox(
                "Valor considerado evento/resultado positivo",
                options=unique_values,
                key="outcome_event_value",
                on_change=invalidate_outcome_confirmation,
            )
            event_rate = float((outcome_series == outcome_event_value).mean())
            st.caption(f"Frecuencia observada del evento en el dataset: {event_rate:.1%}.")
        elif outcome_type == "continuous":
            numeric_outcome = pd.to_numeric(outcome_series, errors="coerce").dropna()
            if numeric_outcome.empty:
                outcome_column_valid = False
                st.error("La columna seleccionada no contiene valores numéricos utilizables para un outcome continuo.")
            else:
                stats_c1, stats_c2, stats_c3 = st.columns(3)
                stats_c1.metric("Observaciones válidas", f"{len(numeric_outcome):,}")
                stats_c2.metric("Media", f"{numeric_outcome.mean():.3f}")
                stats_c3.metric("DE", f"{numeric_outcome.std(ddof=1):.3f}" if len(numeric_outcome) > 1 else "—")
        elif outcome_type == "time_to_event":
            st.info(
                "Para el análisis final de supervivencia suelen necesitarse una variable de tiempo y otra de evento. Aquí se confirma la columna principal; documente la segunda variable en la definición operacional."
            )
else:
    st.info("No se cargó un dataset. La confirmación del outcome se realizará únicamente contra el protocolo y la definición clínica.")

structured_outcomes = st.session_state.get("protocol_outcomes", [])
if structured_outcomes:
    with st.expander("Outcomes estructurados detectados en el protocolo", expanded=False):
        st.dataframe(pd.DataFrame(structured_outcomes), hide_index=True, use_container_width=True)
    coprimary_count = sum(
        1 for item in structured_outcomes if item.get("es_primario") or item.get("es_coprimario")
    )
    if coprimary_count > 1:
        st.warning(
            "Se detectaron varios outcomes primarios/coprimarios. La app calculará escenarios adicionales cuando cada uno contenga efecto, variabilidad y diseño suficientes, y adoptará el mayor N como recomendación global."
        )

confirmation_ready = bool(outcome_primary.strip()) and outcome_type != "unknown"
if dataset is not None:
    confirmation_ready = confirmation_ready and bool(selected_outcome_column) and outcome_column_valid

outcome_confirmed = st.checkbox(
    "Confirmo que este es el outcome primario preespecificado y que su tipo, momento y vinculación con los datos son correctos.",
    key="outcome_confirmed",
    disabled=not confirmation_ready,
)
if not confirmation_ready:
    st.error(
        "Complete el nombre y tipo del outcome primario"
        + (" y seleccione su columna en el dataset" if dataset is not None else "")
        + "."
    )

outcome_metadata: dict[str, Any] = {
    "outcome_primario": outcome_primary.strip(),
    "definicion_outcome": outcome_definition.strip(),
    "tipo_outcome": outcome_type,
    "tipo_outcome_label": OUTCOME_TYPE_LABELS.get(outcome_type, outcome_type),
    "unidad_outcome": outcome_unit.strip(),
    "momento_evaluacion": outcome_timepoint.strip(),
    "outcome_secundarios": outcome_secondary.strip(),
    "columna_dataset_outcome": selected_outcome_column,
    "confianza_columna": outcome_column_score,
    "valor_evento": outcome_event_value if outcome_type in {"binary", "paired_binary", "diagnostic"} else "",
    "confirmado": bool(outcome_confirmed),
}

with st.expander("⚙️ Parámetros estadísticos generales", expanded=True):
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        alpha = st.number_input("Error alfa", min_value=0.001, max_value=0.20, step=0.005, format="%.3f", key="alpha")
    with c2:
        power = st.number_input("Poder estadístico (0–1)", min_value=0.50, max_value=0.99, step=0.01, format="%.2f", key="power")
        if power < 0.80:
            st.caption("⚠️ Inferior al mínimo convencional de 80%.")
        elif math.isclose(power, 0.80, abs_tol=1e-9):
            st.caption("80%: mínimo aceptado habitualmente.")
    with c3:
        sided = st.radio("Tipo de contraste", ["Bilateral", "Unilateral"], horizontal=True, key="sided")
    with c4:
        loss_percent = st.slider("Pérdidas esperadas (%)", 0, 50, step=1, key="loss_percent")
        loss_rate = loss_percent / 100.0

# Parámetros específicos
st.subheader("Parámetros del diseño")
params: dict[str, Any] = {
    "design_code": design_code,
    "alpha": alpha,
    "power": power,
    "sided": sided,
    "loss_rate": loss_rate,
    "outcome": outcome_metadata,
}

if design_code == "mean_estimation":
    c1, c2, c3 = st.columns(3)
    with c1:
        params["sd"] = st.number_input("Desviación estándar esperada", min_value=0.0001, key="sd")
    with c2:
        params["precision"] = st.number_input("Error máximo aceptable (±d)", min_value=0.0001, key="precision_mean")
    with c3:
        params["finite_population"] = st.checkbox("Aplicar población finita", key="finite_population")
        if params["finite_population"]:
            params["population"] = st.number_input("Tamaño de la población", min_value=2, step=1, key="population")

elif design_code == "proportion_estimation":
    c1, c2, c3 = st.columns(3)
    with c1:
        params["p"] = st.number_input("Proporción esperada", min_value=0.001, max_value=0.999, step=0.01, key="p")
        st.caption("Use 0,50 cuando sea desconocida para un cálculo conservador.")
    with c2:
        params["precision"] = st.number_input("Error máximo aceptable (±d)", min_value=0.001, max_value=0.50, step=0.005, key="precision_prop")
    with c3:
        params["finite_population"] = st.checkbox("Aplicar población finita", key="finite_population")
        if params["finite_population"]:
            params["population"] = st.number_input("Tamaño de la población", min_value=2, step=1, key="population")

elif design_code in {"means_independent", "superiority_continuous"}:
    c1, c2 = st.columns(2)
    with c1:
        params["delta"] = st.number_input("Diferencia mínima clínicamente importante", min_value=0.0001, key="delta")
    with c2:
        params["sd"] = st.number_input("Desviación estándar común", min_value=0.0001, key="sd")

elif design_code == "means_paired":
    c1, c2 = st.columns(2)
    with c1:
        params["delta"] = st.number_input("Diferencia media esperada", min_value=0.0001, key="delta")
    with c2:
        params["sd_diff"] = st.number_input("DE de las diferencias intraindividuo", min_value=0.0001, key="sd_diff")

elif design_code in {"proportions_independent", "superiority_binary"}:
    c1, c2, c3 = st.columns(3)
    with c1:
        params["p1"] = st.number_input("Proporción grupo 1", min_value=0.001, max_value=0.999, step=0.01, key="p1")
    with c2:
        params["p2"] = st.number_input("Proporción grupo 2", min_value=0.001, max_value=0.999, step=0.01, key="p2")
    with c3:
        params["yates"] = st.checkbox("Corrección de continuidad", key="yates")

elif design_code == "mcnemar":
    c1, c2, c3 = st.columns(3)
    with c1:
        params["p01"] = st.number_input("Discordantes 0→1", min_value=0.001, max_value=0.999, step=0.01, key="p01")
    with c2:
        params["p10"] = st.number_input("Discordantes 1→0", min_value=0.001, max_value=0.999, step=0.01, key="p10")
    with c3:
        params["yates"] = st.checkbox("Corrección de continuidad", key="yates")

elif design_code in {"noninferiority_continuous", "equivalence_continuous"}:
    c1, c2, c3 = st.columns(3)
    with c1:
        params["sd"] = st.number_input("Desviación estándar común", min_value=0.0001, key="sd")
    with c2:
        params["margin"] = st.number_input("Margen clínico", min_value=0.0001, key="margin_cont")
    with c3:
        params["expected_diff"] = st.number_input("Diferencia esperada T−C", key="expected_diff")
    params["sided"] = "Unilateral"

elif design_code == "noninferiority_binary":
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        params["p1"] = st.number_input("Proporción tratamiento", min_value=0.001, max_value=0.999, step=0.01, key="p1_ni")
    with c2:
        params["p2"] = st.number_input("Proporción control", min_value=0.001, max_value=0.999, step=0.01, key="p2_ni")
    with c3:
        params["margin"] = st.number_input("Margen absoluto", min_value=0.001, max_value=0.50, step=0.01, key="margin_bin")
    with c4:
        params["yates"] = st.checkbox("Corrección de continuidad", key="yates")
    params["sided"] = "Unilateral"

elif design_code == "equivalence_binary":
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        params["p1"] = st.number_input("Proporción tratamiento", min_value=0.001, max_value=0.999, step=0.01, key="p1_eq")
    with c2:
        params["p2"] = st.number_input("Proporción control", min_value=0.001, max_value=0.999, step=0.01, key="p2_eq")
    with c3:
        params["margin"] = st.number_input("Margen absoluto", min_value=0.001, max_value=0.50, step=0.01, key="margin_bin")
    with c4:
        params["yates"] = st.checkbox("Corrección de continuidad", key="yates")
    params["sided"] = "Unilateral"

elif design_code in {"odds_ratio", "risk_ratio"}:
    c1, c2, c3 = st.columns(3)
    with c1:
        params["p_control"] = st.number_input("Proporción/riesgo basal en control", min_value=0.001, max_value=0.999, step=0.01, key="p_control")
    with c2:
        label = "Odds Ratio mínimo" if design_code == "odds_ratio" else "Riesgo Relativo mínimo"
        params["effect_ratio"] = st.number_input(label, min_value=0.01, step=0.05, key="effect_ratio")
    with c3:
        params["yates"] = st.checkbox("Corrección de continuidad", key="yates")

elif design_code == "diagnostic_accuracy":
    c1, c2, c3 = st.columns(3)
    with c1:
        params["sensitivity"] = st.number_input("Sensibilidad esperada", min_value=0.001, max_value=0.999, step=0.01, key="sensitivity")
        params["precision_se"] = st.number_input("Precisión para sensibilidad (±d)", min_value=0.001, max_value=0.30, step=0.005, key="precision_se")
    with c2:
        params["specificity"] = st.number_input("Especificidad esperada", min_value=0.001, max_value=0.999, step=0.01, key="specificity")
        params["precision_sp"] = st.number_input("Precisión para especificidad (±d)", min_value=0.001, max_value=0.30, step=0.005, key="precision_sp")
    with c3:
        params["prevalence"] = st.number_input("Prevalencia de la enfermedad", min_value=0.001, max_value=0.999, step=0.01, key="prevalence")
        st.caption("Se calcularán NSe y NSp por separado y se recomendará el mayor.")

elif design_code == "survival":
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        params["hazard_ratio"] = st.number_input("Hazard Ratio objetivo", min_value=0.01, step=0.05, key="hazard_ratio")
    with c2:
        params["mortality_control"] = st.number_input("Mortalidad acumulada control", min_value=0.001, max_value=0.999, step=0.01, key="mortality_control")
    with c3:
        params["mortality_treatment"] = st.number_input("Mortalidad acumulada tratamiento", min_value=0.001, max_value=0.999, step=0.01, key="mortality_treatment")
    with c4:
        params["allocation_treatment"] = st.number_input("Fracción asignada a tratamiento", min_value=0.05, max_value=0.95, step=0.05, key="allocation_treatment")

st.markdown(
    """
    <div class="technical-note"><b>Nota clínica:</b> la diferencia mínima clínicamente importante debe definirse antes del análisis. Un efecto puede ser estadísticamente significativo y, aun así, ser clínicamente irrelevante.</div>
    """,
    unsafe_allow_html=True,
)

# Asistente de estimación desde el outcome confirmado
if dataset is not None and selected_outcome_column:
    with st.expander("🧪 Estimar parámetros desde el outcome confirmado", expanded=False):
        actual_column = next(column for column in dataset.columns if str(column) == selected_outcome_column)
        outcome_values = dataset[actual_column].dropna()
        st.dataframe(dataset[[actual_column]].head(100), use_container_width=True, height=260)

        if outcome_type == "continuous":
            clean = pd.to_numeric(outcome_values, errors="coerce").dropna()
            if len(clean) > 1:
                estimated_sd = float(clean.std(ddof=1))
                st.write(
                    {
                        "n": int(clean.size),
                        "media": round(float(clean.mean()), 4),
                        "DE": round(estimated_sd, 4),
                        "mínimo": round(float(clean.min()), 4),
                        "máximo": round(float(clean.max()), 4),
                    }
                )
                if st.button("Usar la DE del outcome en el cálculo", key="use_outcome_sd"):
                    st.session_state.sd = estimated_sd
                    st.session_state.sd_diff = estimated_sd
                    st.rerun()
            else:
                st.info("No hay suficientes observaciones numéricas para estimar la desviación estándar.")

        elif outcome_type in {"binary", "paired_binary", "diagnostic"}:
            if len(outcome_values):
                estimated_p = float((outcome_values == outcome_event_value).mean())
                st.metric("Proporción observada del evento", f"{estimated_p:.3f}")
                if st.button("Usar como proporción/prevalencia", key="use_outcome_proportion"):
                    bounded = min(max(estimated_p, 0.001), 0.999)
                    st.session_state.p = bounded
                    st.session_state.prevalence = bounded
                    st.session_state.p_control = bounded
                    st.rerun()
        else:
            st.info(
                "La estimación automática desde el dataset está disponible para outcomes continuos y binarios. Para supervivencia se requieren tiempo, evento y supuestos de seguimiento."
            )

# Cálculo y dashboard
if not outcome_confirmed:
    st.warning("El cálculo está bloqueado hasta confirmar el outcome primario.")
    st.stop()
if not design_compatible_with_outcome(design_code, outcome_type):
    st.error("El tipo de outcome confirmado no es compatible con el diseño estadístico seleccionado.")
    st.stop()
if dataset is not None and outcome_column_score < 0.45:
    st.warning(
        "La correspondencia automática entre el outcome y la columna seleccionada es baja. El cálculo continúa porque fue confirmada manualmente; documente esta decisión."
    )

try:
    result = calculate_sample_size(params)
except Exception as exc:
    st.error(f"No se pudo completar el cálculo: {exc}")
    st.stop()

coprimary_df, coprimary_warnings = calculate_structured_outcome_scenarios(
    st.session_state.get("protocol_outcomes", []),
    params,
)
primary_row = pd.DataFrame(
    [
        {
            "Outcome": outcome_primary,
            "Rol": "Primario confirmado",
            "Tipo": OUTCOME_TYPE_LABELS.get(outcome_type, outcome_type),
            "Diseño": DESIGNS[design_code],
            "N total": result.n_final_total,
            "Estado": "Calculado y confirmado",
        }
    ]
)
if coprimary_df.empty:
    coprimary_df = primary_row
elif normalize_token(outcome_primary) not in {normalize_token(value) for value in coprimary_df["Outcome"].astype(str)}:
    coprimary_df = pd.concat([primary_row, coprimary_df], ignore_index=True)
numeric_outcome_sizes = pd.to_numeric(coprimary_df["N total"], errors="coerce").dropna()
recommended_total = int(max([result.n_final_total] + numeric_outcome_sizes.astype(int).tolist()))

st.subheader("Dashboard de resumen")
metric_cols = st.columns(4)
metric_cols[0].markdown(
    f'<div class="metric-card"><div class="metric-label">N total recomendado global</div><div class="metric-value">{recommended_total:,}</div><div class="metric-note">Mayor requerimiento entre outcomes · incluye pérdidas</div></div>',
    unsafe_allow_html=True,
)
metric_cols[1].markdown(
    f'<div class="metric-card"><div class="metric-label">N inicial</div><div class="metric-value">{math.ceil(result.n_raw_total):,}</div><div class="metric-note">Antes del ajuste por pérdidas</div></div>',
    unsafe_allow_html=True,
)
first_group = next(iter(result.groups_final.items())) if result.groups_final else ("Grupo", result.n_final_total)
metric_cols[2].markdown(
    f'<div class="metric-card"><div class="metric-label">{first_group[0]}</div><div class="metric-value">{first_group[1]:,}</div><div class="metric-note">Asignación calculada</div></div>',
    unsafe_allow_html=True,
)
if len(result.groups_final) > 1:
    second_group = list(result.groups_final.items())[1]
    metric_cols[3].markdown(
        f'<div class="metric-card"><div class="metric-label">{second_group[0]}</div><div class="metric-value">{second_group[1]:,}</div><div class="metric-note">Asignación calculada</div></div>',
        unsafe_allow_html=True,
    )
else:
    metric_cols[3].markdown(
        f'<div class="metric-card"><div class="metric-label">Poder / criterio</div><div class="metric-value">{power:.0%}</div><div class="metric-note">{"Usado" if result.uses_power else "No aplicable: diseño por precisión"}</div></div>',
        unsafe_allow_html=True,
    )

if result.extra.get("required_events"):
    st.info(f"Eventos requeridos para supervivencia: **{result.extra['required_events']:,}**.")
if design_code == "diagnostic_accuracy":
    st.info(
        f"Requerimiento por sensibilidad: **{result.extra['n_sensitivity']:,}** · por especificidad: **{result.extra['n_specificity']:,}** · se adopta el mayor."
    )
for warning in result.warnings:
    st.warning(warning)

if len(coprimary_df) > 1 or coprimary_warnings:
    st.markdown("#### Outcomes primarios/coprimarios")
    st.dataframe(coprimary_df, hide_index=True, use_container_width=True)
    if recommended_total > result.n_final_total:
        st.success(
            f"La recomendación global aumenta de {result.n_final_total:,} a **{recommended_total:,}** participantes porque otro outcome primario/coprimario exige una muestra mayor."
        )
    for warning in coprimary_warnings:
        st.warning(f"No se pudo calcular un escenario coprimario: {warning}")

# Curva de potencia o precisión
plot_col, detail_col = st.columns([1.55, 1.0], gap="large")
with plot_col:
    if result.uses_power:
        powers = np.round(np.linspace(0.50, 0.95, 19), 3)
        n_values: list[int] = []
        for pwr in powers:
            try:
                n_values.append(calculate_sample_size(params, power_override=float(pwr)).n_final_total)
            except Exception:
                n_values.append(np.nan)
        fig = go.Figure()
        fig.add_trace(
            go.Scatter(
                x=powers * 100,
                y=n_values,
                mode="lines+markers",
                name="N total",
                hovertemplate="Poder %{x:.0f}%<br>N %{y:,.0f}<extra></extra>",
            )
        )
        fig.add_vline(x=power * 100, line_dash="dash", annotation_text=f"Seleccionado: {power:.0%}")
        fig.update_layout(
            title="Curva de potencia: tamaño de muestra vs. poder",
            xaxis_title="Poder estadístico (%)",
            yaxis_title="N total final",
            height=430,
            margin=dict(l=20, r=20, t=60, b=20),
            hovermode="x unified",
        )
        st.plotly_chart(fig, use_container_width=True)
    else:
        if design_code == "diagnostic_accuracy":
            base_precision = min(params["precision_se"], params["precision_sp"])
            precision_values = np.linspace(max(0.01, base_precision * 0.55), min(0.25, base_precision * 1.60), 20)
            n_values = []
            for d in precision_values:
                pcopy = dict(params)
                pcopy["precision_se"] = float(d)
                pcopy["precision_sp"] = float(d)
                n_values.append(calculate_sample_size(pcopy).n_final_total)
        else:
            base_precision = params["precision"]
            precision_values = np.linspace(max(0.001, base_precision * 0.55), base_precision * 1.60, 20)
            n_values = []
            for d in precision_values:
                pcopy = dict(params)
                pcopy["precision"] = float(d)
                n_values.append(calculate_sample_size(pcopy).n_final_total)
        fig = go.Figure()
        fig.add_trace(
            go.Scatter(
                x=precision_values,
                y=n_values,
                mode="lines+markers",
                name="N total",
                hovertemplate="Precisión ±%{x:.3f}<br>N %{y:,.0f}<extra></extra>",
            )
        )
        fig.add_vline(x=base_precision, line_dash="dash", annotation_text="Precisión seleccionada")
        fig.update_layout(
            title="Curva de precisión: tamaño de muestra vs. error máximo",
            xaxis_title="Error máximo aceptable (±d)",
            yaxis_title="N total final",
            height=430,
            margin=dict(l=20, r=20, t=60, b=20),
        )
        st.plotly_chart(fig, use_container_width=True)

with detail_col:
    st.markdown("#### Justificación matemática")
    st.code(result.formula, language=None)
    st.write(result.method)
    st.markdown("**Supuestos:**")
    st.write(" · ".join(result.assumptions))
    group_df = pd.DataFrame(
        [
            {
                "Componente": name,
                "N sin pérdidas": math.ceil(result.groups_raw.get(name, 0)),
                "N final": value,
            }
            for name, value in result.groups_final.items()
        ]
    )
    st.dataframe(group_df, hide_index=True, use_container_width=True)

# Análisis de sensibilidad y descargas
if result.uses_power:
    scenarios = sorted(set([power, 0.80, 0.90, 0.95]))
    sensitivity_rows = []
    for scenario_power in scenarios:
        scenario_result = calculate_sample_size(params, power_override=scenario_power)
        sensitivity_rows.append(
            {
                "Escenario": f"Poder {scenario_power:.0%}",
                "Poder": f"{scenario_power:.0%}",
                "N total": scenario_result.n_final_total,
                "Variación vs actual": scenario_result.n_final_total - result.n_final_total,
            }
        )
else:
    sensitivity_rows = [
        {
            "Escenario": "Diseño basado en precisión",
            "Poder": "No aplica",
            "N total": result.n_final_total,
            "Variación vs actual": 0,
        },
        {
            "Escenario": "Poder 90%",
            "Poder": "No aplica a esta fórmula",
            "N total": result.n_final_total,
            "Variación vs actual": 0,
        },
    ]
sensitivity_df = pd.DataFrame(sensitivity_rows)

st.subheader("Análisis de sensibilidad")
st.dataframe(sensitivity_df, hide_index=True, use_container_width=True)

word_report = build_word_report(
    result,
    params,
    sensitivity_df,
    dataset_metadata,
    coprimary_df=coprimary_df if len(coprimary_df) > 1 else None,
    recommended_total=recommended_total,
)
consort_docx = build_consort_docx()
consort_pdf = build_consort_pdf()

st.subheader("Descargas")
d1, d2, d3 = st.columns(3)
with d1:
    st.download_button(
        "⬇️ Informe Word editable",
        data=word_report,
        file_name=f"informe_tamano_muestra_{design_code}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
with d2:
    st.download_button(
        "⬇️ Checklist CONSORT 2010 · Word",
        data=consort_docx,
        file_name="checklist_CONSORT_2010_operativo.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
with d3:
    st.download_button(
        "⬇️ Checklist CONSORT 2010 · PDF",
        data=consort_pdf,
        file_name="checklist_CONSORT_2010_operativo.pdf",
        mime="application/pdf",
        use_container_width=True,
    )

with st.expander("Texto CONSORT generado", expanded=False):
    st.write(consort_paragraph(result, params, recommended_total))

st.divider()
st.caption(
    "Herramienta de planificación metodológica. Verifique los supuestos, la fuente del efecto esperado y la adecuación del método al protocolo. "
    "Los estudios regulatorios o de diseño complejo requieren revisión bioestadística independiente."
)
